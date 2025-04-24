import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf

data, fs = sf.read('../SOUND_INTRO/sound1.wav', dtype='float32')

print(data.dtype)
print(data.shape)

# sd.play(data, fs)
# status = sd.wait()


# zadanie 1

sound_L = data[:, 0]
sound_R = data[:, 1]
sound_mix = (data[:, 0] + data[:, 1]) / 2

# sf.write('SOUND_INTRO/sound_L.wav', sound_L, fs)
# sf.write('SOUND_INTRO/sound_R.wav', sound_R, fs)
# sf.write('SOUND_INTRO/sound_mix.wav', sound_mix, fs)


# czesc 3

# x = np.arange(0, data.shape[0]/fs, 1/fs)
#
# plt.subplot(2,1,1)
# plt.plot(x, data[:, 0])
# plt.show()

# widmo

data, fs = sf.read('../SIN/sin_440Hz.wav', dtype=np.int32)
# plt.figure()
# plt.subplot(2,1,1)
# plt.plot(np.arange(0,data.shape[0])/fs,data)
#
# plt.subplot(2,1,2)
# yf = scipy.fftpack.fft(data)
# plt.plot(np.arange(0,fs,1.0*fs/(yf.size)),np.abs(yf))
# plt.show()
#
# fsize = 2**8
# plt.figure()
# plt.subplot(2,1,1)
# plt.plot(np.arange(0,data.shape[0])/fs,data)
# plt.subplot(2,1,2)
# yf = scipy.fftpack.fft(data,fsize)
# plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
# plt.show()



# zadanie 2

from docx import Document
from docx.shared import Inches
from io import BytesIO

def plotAudio(Signal, Fs, TimeMargin=[0,0.02], fsize=2**8):
    fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
    fig.suptitle('Time margin {}'.format(Margin)) # Tytuł wykresu
    fig.tight_layout(pad=1.5) # poprawa czytelności

    axs[0].plot(np.arange(0,Signal.shape[0]/Fs, 1/Fs), data)
    axs[0].axis(xmin=TimeMargin[0], xmax=TimeMargin[1])
    axs[0].set_xlabel('Time (s)')
    axs[0].set_ylabel('Amplitude')
    yf = scipy.fftpack.fft(Signal,fsize)
    x = np.arange(0, fs/2, fs/fsize)
    spectrum = 20*np.log10( np.abs(yf[:fsize//2]))
    axs[1].plot(x, spectrum)
    axs[1].set_xlabel('Frequency (Hz)')
    axs[1].set_ylabel('Amplitude')
    memfile = BytesIO()
    fig.savefig(memfile)

    plt.show()

    max_idx = np.argmax(spectrum)
    max_freq = x[max_idx]
    max_amplitude = spectrum[max_idx]

    return memfile, max_freq, max_amplitude

# data, fs = sf.read('../SIN/sin_8000Hz.wav', dtype=np.int32)
# plotAudio(data, fs, [0, 0.0005])


document = Document()
document.add_heading('Zmień ten tytuł',0) # tworzenie nagłówków druga wartość to poziom nagłówka


files=['sin_60Hz.wav','sin_440Hz.wav','sin_8000Hz.wav']
Margins=[[0,0.02],[0.133,0.155]]
fsize = [2**8, 2**12, 2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file),2)
    for Fsize in fsize:
        for i,Margin in enumerate(Margins):
            document.add_heading(f'Time margin {Margin}, fsize {Fsize}',3) # nagłówek sekcji, mozę być poziom wyżej
            data,fs = sf.read(f'../SIN/{file}')
            print(i, Fsize)
            memfile, max_freq, max_amplitude = plotAudio(data, fs, Margin, Fsize)

            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku

            memfile.close()
            ############################################################
            # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
            document.add_paragraph(f'Maksymalna częstotliwość: {max_freq:.2f} Hz')
            document.add_paragraph(f'Amplituda: {max_amplitude:.2f} dB')
            ############################################################

document.save('report.docx') # zapis do pliku
