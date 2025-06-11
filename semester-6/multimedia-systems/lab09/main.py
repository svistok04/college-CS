from time import time
from io import BytesIO
from docx import Document
from docx.shared import T, Inches
from typing import Optional
from types import FrameType
import cv2
import numpy as np
import matplotlib.pyplot as plt
import os

##############################################################################
######   Konfiguracja       ##################################################
##############################################################################
flag_compress_KeyFrame = False
flag_run_byterun = False
kat='.'                                 # katalog z plikami wideo
plik="clip_4.mp4"                       # nazwa pliku
frames_num = 6                     # ile klatek odtworzyć? <0 - całość
key_frame_counter=2                     # co która klatka ma być kluczowa i nie podlegać kompresji
plot_frames=np.array([30,45])           # automatycznie wyrysuj wykresy
auto_pause_frames=np.array([25])        # automatycznie za pauzuj dla klatki
subsampling="4:1:1"                     # parametry dla chroma subsampling
divisor=2                               # dzielnik przy zapisie różnicy
display_frames=False                     # czy program ma wyświetlać klatki
ROI = [[100,500,100,500]]                   # wyświetlane fragmenty (można podać kilka )


##############################################################################
####     Kompresja i dekompresja    ##########################################
##############################################################################
class data:
    def __init__(self):
        self.Y: np.ndarray | None = None
        self.Cb: np.ndarray | None = None
        self.Cr: np.ndarray | None = None

def Chroma_subsampling(L, subsampling):
    match subsampling:
        case "4:4:4":
            return L
        case "4:2:2":
            return L[:, ::2]
        case "4:4:0":
            return L[::2, :]
        case "4:2:0":
            return L[::2, ::2]
        case "4:1:1":
            return L[:, ::4]
        case "4:1:0":
            return L[::2, ::4]
        case _: 
            return L

def Chroma_resampling(L, subsampling):
    match subsampling:
        case "4:4:4":
            return L
        case "4:2:2":
            return np.repeat(L, 2, axis=1)
        case "4:4:0":
            return np.repeat(L, 2, axis=0)
        case "4:2:0":
            return np.repeat(np.repeat(L, 2, axis=0), 2, axis=1)
        case "4:1:1":
            return np.repeat(L, 4, axis=1)
        case "4:1:0":
            return np.repeat(np.repeat(L, 2, axis=0), 4, axis=1)
        case _:
            return L

def frame_image_to_class(frame,subsampling):
    Frame_class = data()
    Frame_class.Y=frame[:,:,0].astype(int)
    Frame_class.Cb=Chroma_subsampling(frame[:,:,2].astype(int),subsampling)
    Frame_class.Cr=Chroma_subsampling(frame[:,:,1].astype(int),subsampling)
    return Frame_class


def frame_layers_to_image(Y,Cr,Cb,subsampling):  
    Cb=Chroma_resampling(Cb,subsampling)
    Cr=Chroma_resampling(Cr,subsampling)
    return np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)

def compress_KeyFrame(Frame_class):
    KeyFrame = data()
    if flag_compress_KeyFrame and flag_run_byterun:
        KeyFrame.Y = byterun_encode(Frame_class.Y.astype(np.uint8))
        KeyFrame.Cb = byterun_encode(Frame_class.Cb.astype(np.uint8))
        KeyFrame.Cr = byterun_encode(Frame_class.Cr.astype(np.uint8))
    else:
        KeyFrame.Y = Frame_class.Y
        KeyFrame.Cb =Frame_class.Cb
        KeyFrame.Cr =Frame_class.Cr
    return KeyFrame

def decompress_KeyFrame(KeyFrame):
    if flag_compress_KeyFrame and flag_run_byterun:
        Y = byterun_decode(KeyFrame.Y)
        Cb = byterun_decode(KeyFrame.Cb)
        Cr = byterun_decode(KeyFrame.Cr)
    else: 
        Y=KeyFrame.Y
        Cb=KeyFrame.Cb
        Cr=KeyFrame.Cr
    return frame_layers_to_image(Y, Cr, Cb, subsampling)

def compress_not_KeyFrame(Frame_class, KeyFrame, divisor=1):
    Compress_data = data()
    if flag_compress_KeyFrame and flag_run_byterun:
        Compress_data.Y  = byterun_encode(((Frame_class.Y  - byterun_decode(KeyFrame.Y))  // divisor).astype(np.int8))
        Compress_data.Cb = byterun_encode(((Frame_class.Cb - byterun_decode(KeyFrame.Cb)) // divisor).astype(np.int8))
        Compress_data.Cr = byterun_encode(((Frame_class.Cr - byterun_decode(KeyFrame.Cr)) // divisor).astype(np.int8))
    elif not flag_compress_KeyFrame and flag_run_byterun:
        Compress_data.Y  = byterun_encode(((Frame_class.Y  - (KeyFrame.Y))  // divisor).astype(np.int8))
        Compress_data.Cb = byterun_encode(((Frame_class.Cb - (KeyFrame.Cb)) // divisor).astype(np.int8))
        Compress_data.Cr = byterun_encode(((Frame_class.Cr - (KeyFrame.Cr)) // divisor).astype(np.int8))
    elif not flag_run_byterun:
        Compress_data.Y  = Frame_class.Y - KeyFrame.Y // divisor
        Compress_data.Cb = Frame_class.Cb- KeyFrame.Cb  // divisor
        Compress_data.Cr = Frame_class.Cr- KeyFrame.Cr  // divisor
    return Compress_data

def decompress_not_KeyFrame(Compress_data, KeyFrame, divisor=1):
    if flag_run_byterun:
        diff_Y  = byterun_decode(Compress_data.Y).astype(np.int16)
        diff_Cb = byterun_decode(Compress_data.Cb).astype(np.int16)
        diff_Cr = byterun_decode(Compress_data.Cr).astype(np.int16)
        if flag_compress_KeyFrame:
            Y  = byterun_decode(KeyFrame.Y)  + diff_Y  * divisor
            Cb = byterun_decode(KeyFrame.Cb) + diff_Cb * divisor
            Cr = byterun_decode(KeyFrame.Cr) + diff_Cr * divisor
        else:
            Y  = (KeyFrame.Y)  + diff_Y  * divisor
            Cb = (KeyFrame.Cb) + diff_Cb * divisor
            Cr = (KeyFrame.Cr) + diff_Cr * divisor
    else:
            Y  = (KeyFrame.Y)  + Compress_data.Y  * divisor
            Cb = (KeyFrame.Cb) + Compress_data.Cb * divisor
            Cr = (KeyFrame.Cr) + Compress_data.Cr * divisor
    return frame_layers_to_image(Y, Cr, Cb, subsampling)

def plot_diffrence(ReferenceFrame_RGB, DecompressedFrame_RGB, ROI):
    fig, axs = plt.subplots(1, 3 , sharey=True)
    fig.set_size_inches(16,5)

    axs[0].imshow(ReferenceFrame_RGB[ROI[0]:ROI[1], ROI[2]:ROI[3]])
    axs[2].imshow(DecompressedFrame_RGB[ROI[0]:ROI[1], ROI[2]:ROI[3]])

    diff = ReferenceFrame_RGB[ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float) - DecompressedFrame_RGB[ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float)
    # axs[1].imshow(diff, vmin=np.min(diff),vmax=np.max(diff))
    axs[1].imshow(np.abs(diff).mean(axis=2), cmap='gray')
    axs[0].set_title("Original")
    axs[1].set_title("Difference")
    axs[2].set_title("Reconstructed")
    for ax in axs:
        ax.axis("off")
    
    plt.tight_layout()
    if not flag_run_byterun:
        fname = f"no_byterun_frames/diff_sub_{subsampling}_div{divisor}_.png"
    else:
        fname = f"byterun_frames/diff_sub_{subsampling}_div{divisor}_.png"
    plt.savefig(fname)
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    plt.close()
    return buf

def plot_compression_result(compression_information, subsampling, divisor):
    plt.figure()
    plt.plot(np.arange(0,frames_num),compression_information[0,:]*100)
    plt.plot(np.arange(0,frames_num),compression_information[1,:]*100)
    plt.plot(np.arange(0,frames_num),compression_information[2,:]*100)
    # print(f'mean compress 1 {sum(compression_information[0, :]) / float(len(compression_information[0, :]))}, '
        # f'2 {sum(compression_information[1, :]) / float(len(compression_information[1, :]))}, '
        # f'3 {sum(compression_information[2, :]) / float(len(compression_information[2, :]))}')
    plt.title("File:{}, subsampling={}, divider={}, KeyFrame={} ".format(plik,subsampling,divisor,key_frame_counter))
    plt.legend(['Y', 'Cb', 'Cr'])
    if not flag_run_byterun:
        plt.savefig(f"no_byterun/plot_{plik}_{subsampling}_d{divisor}_k{key_frame_counter}.png")
    else:
        plt.savefig(f"byterun/plot_{plik}_{subsampling}_d{divisor}_k{key_frame_counter}.png")
    buf = BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return buf

def byterun_encode(data):
    shape = data.shape
    if len(shape) > 1:
        data = data.flatten()

    length = len(data)
    result = np.zeros(length * 2, dtype=np.uint8)
    out_idx = 0
    i = 0

    while i < length:
        run_start = i
        run_value = data[i]
        run_length = 1
        while i + run_length < length and data[i + run_length] == run_value and run_length < 127:
            # print(run_value)
            run_length += 1

        if run_length > 1:
            # print(f'repeated run length {run_length}')
            result[out_idx] = (-run_length + 256) % 256
            result[out_idx + 1] = run_value
            out_idx += 2
            i += run_length
        else:
            literal_start = i
            literal_length = 1
            while i + literal_length < length - 1 and (data[i + literal_length] != data[i + literal_length + 1] or literal_length == 1) and literal_length < 127:
                literal_length += 1
            result[out_idx] = literal_length
            result[out_idx + 1:out_idx + literal_length + 1] = data[i:i + literal_length]
            out_idx += 1 + literal_length
            i += literal_length

    return np.concatenate((
        np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.uint8),
                result[:out_idx].view(np.uint8)
    ))
    # return np.concatenate((
    #     np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.int8),
    #     result[:out_idx]
    # ))

def byterun_decode(encoded):
    ndim = encoded[:4].view(np.uint32)[0]
    shape = encoded[4:4 + 4 * ndim].view(np.uint32)
    data_start = 4 + 4 * ndim

    data = encoded[data_start:].view(np.int8)
    result = []

    i = 0
    while i < len(data):
        count = data[i]
        if count < 0:
            value = data[i + 1]
            result.extend([value] * (-count))
            i += 2
        else:
            count = int(count)
            result.extend(data[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(result, dtype=np.int8).reshape(tuple(shape))

##############################################################################
####     Głowna pętla programu      ##########################################
##############################################################################
subsampling_list = ["4:4:4", "4:2:2", "4:4:0", "4:2:0", "4:1:1", "4:1:0"]
divisors = [1, 2, 4]

document = Document()
document.add_heading("Video Compression Report", level=0)
document.add_page_break()

flag_run_byterun = False
document.add_heading("No ByteRun", level=1)
for subsampling in subsampling_list:
    document.add_paragraph(f'subsampling = {subsampling}')
    # for divisor in divisors:
    cap = cv2.VideoCapture(os.path.join(kat, plik))

    if frames_num < 0:
        frames_num = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

    if display_frames:
        cv2.namedWindow('Normal Frame')
        cv2.namedWindow('Decompressed Frame')


    compression_information=np.zeros((3,frames_num))
    KeyFrame: data | None = None
    cY: Optional[np.ndarray] = None
    cCb: Optional[np.ndarray] = None
    cCr: Optional[np.ndarray] = None
    start = time()

    for i in range(frames_num):
        print(f"Processing frame {i+1}/{frames_num}", end='\r')
        ret, frame_bgr = cap.read()
        if display_frames:
            cv2.imshow('Normal Frame',frame_bgr)
        frame_YCrCb = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2YCrCb)
        frame_RGB = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB)
        Frame_class = frame_image_to_class(frame_YCrCb, subsampling)
        if (i % key_frame_counter)==0: # pobieranie klatek kluczowych
            KeyFrame = compress_KeyFrame(Frame_class)
            cY=KeyFrame.Y
            cCb=KeyFrame.Cb
            cCr=KeyFrame.Cr
            Decompresed_Frame = decompress_KeyFrame(KeyFrame)
        else: # kompresja
            Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
            cY=Compress_data.Y
            cCb=Compress_data.Cb
            cCr=Compress_data.Cr
            Decompresed_Frame = decompress_not_KeyFrame(Compress_data,  KeyFrame)
            Decompressed_RGB = cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2RGB)

        assert cY is not None and cCb is not None and cCr is not None
        compression_information[0,i]= (frame_YCrCb[:,:,0].size - cY.size)/frame_YCrCb[:,:,0].size
        compression_information[1,i]= (frame_YCrCb[:,:,0].size - cCb.size)/frame_YCrCb[:,:,0].size
        compression_information[2,i]= (frame_YCrCb[:,:,0].size - cCr.size)/frame_YCrCb[:,:,0].size
        if display_frames:
            cv2.imshow('Decompressed Frame',cv2.cvtColor(Decompresed_Frame,cv2.COLOR_YCrCb2BGR))

        # if np.any(plot_frames==i): # rysuj wykresy
        #     for r in ROI:
        #         plotDiffrence(frame_RGB, Decompressed_RGB, r)
        if i == key_frame_counter * 2 - 2:
            for r in ROI:
                buf = plot_diffrence(frame_RGB, Decompressed_RGB, r)
                document.add_picture(buf, width=Inches(5.5))

        if np.any(auto_pause_frames==i):
            cv2.waitKey(-1) #wait until any key is pressed

        k = cv2.waitKey(1) & 0xff

        if k==ord('q'):
            break
        elif k == ord('p'):
            cv2.waitKey(-1) #wait until any key is pressed
    print(f'it took {time() - start} s')

    buf = plot_compression_result(compression_information, subsampling, 1)
    document.add_picture(buf, width=Inches(5.5))
    document.add_page_break()
    # plt.show()
flag_run_byterun = True

subsampling_list = ["4:4:4", "4:2:0", "4:1:0"]
divisors = [2, 4]
for subsampling in subsampling_list:
    document.add_paragraph(f'subsampling = {subsampling}')
    for divisor in divisors:
        document.add_paragraph(f'divisor = {divisor}')
        cap = cv2.VideoCapture(os.path.join(kat, plik))

        if frames_num < 0:
            frames_num = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

        if display_frames:
            cv2.namedWindow('Normal Frame')
            cv2.namedWindow('Decompressed Frame')


        compression_information=np.zeros((3,frames_num))
        KeyFrame: data | None = None
        cY: Optional[np.ndarray] = None
        cCb: Optional[np.ndarray] = None
        cCr: Optional[np.ndarray] = None
        start = time()

        for i in range(frames_num):
            print(f"Processing frame {i+1}/{frames_num}", end='\r')
            ret, frame_bgr = cap.read()
            if display_frames:
                cv2.imshow('Normal Frame',frame_bgr)
            frame_YCrCb = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2YCrCb)
            frame_RGB = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB)
            Frame_class = frame_image_to_class(frame_YCrCb, subsampling)
            if (i % key_frame_counter)==0: # pobieranie klatek kluczowych
                KeyFrame = compress_KeyFrame(Frame_class)
                cY=KeyFrame.Y
                cCb=KeyFrame.Cb
                cCr=KeyFrame.Cr
                Decompresed_Frame = decompress_KeyFrame(KeyFrame)
            else: # kompresja
                Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
                cY=Compress_data.Y
                cCb=Compress_data.Cb
                cCr=Compress_data.Cr
                Decompresed_Frame = decompress_not_KeyFrame(Compress_data,  KeyFrame)
                Decompressed_RGB = cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2RGB)

            assert cY is not None and cCb is not None and cCr is not None
            compression_information[0,i]= (frame_YCrCb[:,:,0].size - cY.size)/frame_YCrCb[:,:,0].size
            compression_information[1,i]= (frame_YCrCb[:,:,0].size - cCb.size)/frame_YCrCb[:,:,0].size
            compression_information[2,i]= (frame_YCrCb[:,:,0].size - cCr.size)/frame_YCrCb[:,:,0].size
            if display_frames:
                cv2.imshow('Decompressed Frame',cv2.cvtColor(Decompresed_Frame,cv2.COLOR_YCrCb2BGR))

            # if np.any(plot_frames==i): # rysuj wykresy
            #     for r in ROI:
            #         plotDiffrence(frame_RGB, Decompressed_RGB, r)
            if i == key_frame_counter * 2 - 2:
                for r in ROI:
                    buf = plot_diffrence(frame_RGB, Decompressed_RGB, r)
                    document.add_picture(buf, width=Inches(5.5))

            if np.any(auto_pause_frames==i):
                cv2.waitKey(-1) #wait until any key is pressed

            k = cv2.waitKey(1) & 0xff

            if k==ord('q'):
                break
            elif k == ord('p'):
                cv2.waitKey(-1) #wait until any key is pressed
        print(f'it took {time() - start} s')

        buf = plot_compression_result(compression_information, subsampling, divisor)
        document.add_picture(buf, width=Inches(5.5))
        document.add_page_break()
        # plt.show()

document.save("video_report.docx")


