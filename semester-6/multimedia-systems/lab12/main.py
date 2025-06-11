import numpy as np
import cv2
from skimage.metrics import structural_similarity as ssim
from docx import Document
from docx.shared import Inches
from io import BytesIO

# IMAGE_WIDTH = 640
# IMAGE_HEIGHT = 480
CANVAS_WIDTH = 640
CANVAS_HEIGHT = 480

def scale_point(pt, out_width, out_height):
    return (
        int(pt[0] * out_width / CANVAS_WIDTH),
        int(pt[1] * out_height / CANVAS_HEIGHT)

    )

# if IMAGE_WIDTH != CANVAS_WIDTH:
#     WIDTH_SCALE = IMAGE_WIDTH / CANVAS_WIDTH
# else:
#     WIDTH_SCALE = 1.0
#
# if IMAGE_HEIGHT != CANVAS_HEIGHT:
#     HEIGHT_SCALE = IMAGE_HEIGHT / CANVAS_HEIGHT
# else:
#     HEIGHT_SCALE = 1.0
#
# PixelCoord = lambda y, x: (int(y * HEIGHT_SCALE), int(x * WIDTH_SCALE))
#
# LAYERS_COUNT = 3
# IMAGE = np.zeros((IMAGE_HEIGHT, IMAGE_WIDTH, LAYERS_COUNT), dtype=np.uint8)

def draw_pixel(img, x, y, color):
    if 0 <= x < img.shape[1] and 0 <= y < img.shape[0]:
        if len(color) == 4:
            alpha = color[3] / 255.0
            img[y, x] = (alpha * np.array(color[:3]) + (1 - alpha) * img[y, x]).astype(np.uint8)
        else:
            img[y, x] = color[:3]

def draw_line(img, p1, p2, color):
    x1, y1 = p1
    x2, y2 = p2

    steep = abs(y2 - y1) > abs(x2 - x1)

    if steep:
        x1, y1 = y1, x1
        x2, y2 = y2, x2

    if x1 > x2:
        x1, y1, x2, y2 = x2, y2, x1, y1

    dx = x2 - x1
    dy = abs(y2 - y1)
    y = y1
    y_step = 1 if y2 > y1 else -1
    p = 2 * dy - dx

    for x in range(x1, x2 + 1):
        if steep:
            if 0 <= y < img.shape[1] and 0 <= x < img.shape[0]:
                draw_pixel(img, x, y, color)

        else:
            if 0 <= x < img.shape[1] and 0 <= y < img.shape[0]:
                draw_pixel(img, x, y, color)
        if p < 0:
            p += 2 * dy
        else:
            y += y_step
            p += 2 * (dy - dx)


def draw_circle(img, center, radius, color):
    x = 0
    y = radius
    p = 5/4 - radius
    cx, cy = center
    cx, cy = int(cx), int(cy)

    while x < y:

        points = [
            (cy + y, cx + x), (cy + y, cx - x), (cy - y, cx + x), (cy - y, cx - x),
            (cy + x, cx + y), (cy + x, cx - y), (cy - x, cx + y), (cy - x, cx - y)
        ]

        for py, px in points:
            if 0 <= px < img.shape[1] and 0 <= py < img.shape[0]:
                draw_pixel(img, px, py, color)

        if p < 0:
            p += 2 * x + 1
        else:
            p += 2 * (x - y) + 1
            y -= 1
        x += 1

def draw_triangle(img, p1, p2, p3, color):
    draw_line(img, p1, p2, color)
    draw_line(img, p2, p3, color)
    draw_line(img, p3, p1, color)

def draw_rectangle(img, p1, p2, color):
    x1, y1 = p1
    x2, y2 = p2
    corners = [(x1, y1), (x2, y1), (x2, y2), (x1, y2)]
    for i in range(4):
        draw_line(img, corners[i], corners[(i + 1) % 4], color)

def draw_polygon(img, points, color):
    n = len(points)
    for i in range(n):
        draw_line(img, points[i], points[(i + 1) % n], color)

def fill_polygon(img, points, color):
    pts = np.array(points, np.int32)
    min_y = np.min(pts[:,1])
    max_y = np.max(pts[:,1])
    for y in range(min_y, max_y+1):
        intersections = []
        n = len(pts)
        for i in range(n):
            x1, y1 = pts[i]
            x2, y2 = pts[(i+1)%n]
            if y1 == y2:
                continue
            if (y >= min(y1, y2)) and (y < max(y1, y2)):
                x = int(x1 + (y - y1) * (x2 - x1) / (y2 - y1))
                intersections.append(x)
        if len(intersections) > 1:
            intersections.sort()
            for x in range(intersections[0], intersections[-1]+1):
                if 0 <= x < img.shape[1] and 0 <= y < img.shape[0]:
                    draw_pixel(img, x, y, color)


def draw_from_dict(image_shape, image_data):
    out_height, out_width = image_shape[0], image_shape[1]
    bg = image_data.get('Canvas', {}).get('background_color', [0,0,0])
    img = np.ones(image_shape, dtype=np.uint8) * np.array(bg, dtype=np.uint8)

    for shape in sorted(image_data['Shapes'], key=lambda s: s.get('Z_layer', 0)):
        t = shape['type']
        color = shape['color']
        fill = shape.get("fill", False)

        match t:
            case 'circle':
                center = scale_point(shape['center'], out_width, out_height)
                radius = int(shape['radius'] * (out_width / CANVAS_WIDTH))
                if fill and len(color) == 3:
                    cv2.circle(img, center, radius, color, thickness=-1)
                else:
                    draw_circle(img, center, radius, color)
            case 'line':
                p1 = scale_point(shape['p1'], out_width, out_height)
                p2 = scale_point(shape['p2'], out_width, out_height)
                draw_line(img, p1, p2, color)
            case 'triangle':
                pts = [scale_point(shape['p1'], out_width, out_height),
                       scale_point(shape['p2'], out_width, out_height),
                       scale_point(shape['p3'], out_width, out_height)]
                if fill:
                    fill_polygon(img, pts, color)
                else:
                    draw_triangle(img, *pts, color)
            case 'rectangle':
                x1, y1 = scale_point(shape['p1'], out_width, out_height)
                x2, y2 = scale_point(shape['p2'], out_width, out_height)
                pts = np.array([[x1, y1], [x2, y1], [x2, y2], [x1, y2]], np.int32)
                if fill:
                    fill_polygon(img, pts, color)
                else:
                    draw_rectangle(img, (x1, y1), (x2, y2), color)
            case 'polygon' | 'freeform':
                pts = [scale_point(p, out_width, out_height) for p in shape['points']]
                if fill:
                    fill_polygon(img, pts, color)
                else:
                    draw_polygon(img, pts, color)
    return img



def compare_images(img1, img2):
    mse = np.mean((img1.astype("float") - img2.astype("float")) ** 2)
    ssim_val = ssim(img1, img2, channel_axis=2)
    return mse, ssim_val



Example = {
    "Canvas": {
        "width": 640,
        "height": 480,
        "background_color": [10, 50, 10]
    },
    "Shapes": [
        # 1. Duży okręg częściowo przesłonięty trójkątem,
        {
            "type": "circle",
            "Z_layer": 0,
            "center": [180, 370],
            "radius": 90,
            "color": [100, 200, 255],
            "fill": True
        },
        {
            "type": "triangle",
            "Z_layer": 1,
            "p1": [120, 320],
            "p2": [240, 320],
            "p3": [180, 220],
            "color": [200, 0, 80, 200],
            "fill": True
        },

        # 2. Prostokąt z dwoma kwadratami w środku w innym kolorze,
        {
            "type": "rectangle",
            "Z_layer": 0,
            "p1": [350, 50],
            "p2": [550, 200],
            "color": [100, 70, 70],
            "fill": True
        },
        {
            "type": "rectangle",
            "Z_layer": 1,
            "p1": [375, 75],
            "p2": [425, 125],
            "color": [255, 255, 0],
            "fill": True
        },
        {
            "type": "rectangle",
            "Z_layer": 1,
            "p1": [475, 125],
            "p2": [525, 175],
            "color": [0, 255, 0],
            "fill": True
        },

        # 3. Wielokąt przedstawiający literę L lub V,
        {
            "type": "polygon",
            "Z_layer": 1,
            "points": [[340, 310], [340, 420], [410, 420], [410, 390], [370, 390], [370, 310]],
            "color": [0, 0, 0],
            "fill": True
        },

        # 4. Żółty okrąg wysunięty w połowie za brązowego prostokąta,
        {
            "type": "rectangle",
            "Z_layer": 2,
            "p1": [100, 80],
            "p2": [220, 130],
            "color": [139, 69, 19, 245],
            "fill": True
        },
        {
            "type": "circle",
            "Z_layer": 1,
            "center": [160, 80],
            "radius": 40,
            "color": [255, 255, 0],
            "fill": True
        },

        # 5. 3x kwadrat różowy, biały, szary żaden z kwadratów nie na chodzi na inny
        {
            "type": "rectangle",
            "Z_layer": 2,
            "p1": [60, 150],
            "p2": [110, 200],
            "color": [255, 182, 193],
            "fill": True
        },
        {
            "type": "rectangle",
            "Z_layer": 2,
            "p1": [120, 150],
            "p2": [170, 200],
            "color": [255, 255, 255],
            "fill": True
        },
        {
            "type": "rectangle",
            "Z_layer": 2,
            "p1": [180, 150],
            "p2": [230, 200],
            "color": [192, 192, 192],
            "fill": True
        }
    ]
}

document = Document()
document.add_heading("Vector Data Rasterization Report", 0)


sizes = [(640, 480), (320, 240), (1280, 960), (800, 600), (400, 300)]
images = []
for w, h in sizes:
    img = draw_from_dict((h, w, 3), Example)
    images.append((img, w, h))

document.add_heading("Result Images", 1)
ref = images[0][0]
metrics = []
for i, (img, w, h) in enumerate(images):
    img_resized = cv2.resize(img, (CANVAS_WIDTH, CANVAS_HEIGHT))
    _, buf = cv2.imencode('.png', cv2.cvtColor(img_resized, cv2.COLOR_RGB2BGR))
    img_stream = BytesIO(buf.tobytes())
    document.add_paragraph(f"Image {i+1}: Original Size {w}x{h}, Scaled to {CANVAS_WIDTH}x{CANVAS_HEIGHT}")
    document.add_picture(img_stream, width=Inches(4.0))
    mse, ssim_val = compare_images(ref, img_resized)
    metrics.append((w, h, mse, ssim_val))

document.add_heading("Quality Metrics Table for Scaling", 1)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Original Size"
hdr_cells[1].text = "MSE"
hdr_cells[2].text = "SSIM"
for w, h, mse, ssim_val in metrics:
    row_cells = table.add_row().cells
    row_cells[0].text = f"{w}x{h}"
    row_cells[1].text = f"{mse:.2f}"
    row_cells[2].text = f"{ssim_val:.4f}"

document.save("rasterization_lab_report.docx")
