import numpy as np
import matplotlib.pyplot as plt
import cv2


img1 = plt.imread('IMG_INTRO/A1.png')
print('plt, float32')
print(img1.dtype)
print(img1.shape)
print(np.min(img1), np.max(img1))


img2 = cv2.imread('IMG_INTRO/A1.png')
print('cv, uint8')
print(img2.dtype)
print(img2.shape)
print(np.min(img2), np.max(img2))

# zadanie 1

def imgToUint8(img):
    if np.issubdtype(img.dtype, np.unsignedinteger):
        return img
    img *= 255.0
    return img.astype(np.uint8)

def imgToFloat(img):
    if np.issubdtype(img.dtype, np.floating):
        return img

    return img / 255.0

img2 = imgToFloat(img2)
img1 = imgToUint8(img1)

print('imgToUint8')
print(img1.dtype)
print(img1.shape)
print(np.min(img1), np.max(img1))

print('imgToFloat')
print(img2.dtype)
print(img2.shape)
print(np.min(img2), np.max(img2))


# zadanie 2
img = cv2.imread('IMG_INTRO/B02.jpg')
img_RGB = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
# plt.imshow(img_RGB)
# plt.show()

R = img_RGB[:,:,0]
# plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=255)
# plt.show()


R = img_RGB[:,:,0]
G = img_RGB[:,:,1]
B = img_RGB[:,:,2]

Y1 = 0.299 * R + 0.587 * G + 0.114 * B
# plt.imshow(Y1)
# plt.show()

Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B


img = cv2.imread('IMG_INTRO/A4.jpg')

from docx import Document
from docx.shared import Inches
from io import BytesIO

def plotImgFromCV(img_RGB):
    fig, axs = plt.subplots(3, 3, figsize=(10, 10))
    fig.tight_layout(pad=1.5)
    # img_RGB = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    R = img_RGB[:,:,0]
    G = img_RGB[:,:,1]
    B = img_RGB[:,:,2]
    axs[0,0].imshow(img_RGB)
    axs[0,0].title.set_text('Original')
    Y1 = 0.299 * R + 0.587 * G + 0.114 * B
    axs[0,1].imshow(Y1, cmap=plt.cm.gray, vmin=0, vmax=255)
    axs[0,1].title.set_text('Y1')
    Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B
    axs[0,2].imshow(Y2, cmap=plt.cm.gray, vmin=0, vmax=255)
    axs[0,2].title.set_text('Y2')
    axs[1,0].imshow(R, cmap=plt.cm.gray, vmin=0, vmax=255)
    axs[1,0].title.set_text('R')
    axs[1,1].imshow(G, cmap=plt.cm.gray, vmin=0, vmax=255)
    axs[1,1].title.set_text('G')
    axs[1,2].imshow(B, cmap=plt.cm.gray, vmin=0, vmax=255)
    axs[1,2].title.set_text('B')
    R = img_RGB.copy()
    R[:,:,1] = 0
    R[:,:,2] = 0
    axs[2,0].imshow(R)
    axs[2,0].set_title('R', color='red')
    G = img_RGB.copy()
    G[:,:,0] = 0
    G[:,:,2] = 0
    axs[2,1].imshow(G)
    axs[2,1].set_title('G', color='green')
    B = img_RGB.copy()
    B[:,:,0] = 0
    B[:,:,1] = 0
    axs[2,2].imshow(B)
    axs[2,2].set_title('B', color='blue')
    # plt.show()

    memfile = BytesIO()
    fig.savefig(memfile)
    return memfile


plotImgFromCV(img)

img = cv2.imread('IMG_INTRO/B01.png')
img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
fragment = img[10:210, 100:300].copy()
# plt.imshow(fragment)
# plt.show()


import pandas as pd

df = pd.DataFrame()
df = pd.DataFrame(data=[{'Filename':['B01.png'],'Grayscale':[False],
                        'Fragments':[[[20,20,60,60],[100,100,140,140], [60,60,100,100], [140,140,200,200]]]
                        },
                        {'Filename':['B02.jpg'],'Grayscale':[False],
                            'Fragments':[[[20,20,60,60],[100,100,140,140], [50, 0, 500, 200]]]
                         },
                        {'Filename':['A4.jpg'],'Grayscale':[False],
                         'Fragments':[[[2,2,6,6],[10,10,14,14], [5, 80, 50, 100]]]}])

print(df)
# W jaki sposób przetworzyć taką listę? Poniżej przykład:

document = Document()
document.add_heading('Zmień ten tytuł',0)

for index, row in df.iterrows():
    img = plt.imread(f'IMG_INTRO/{row['Filename'][0]}')
    document.add_heading(f'Plik - {row['Filename']}',2)
    if row['Grayscale']:
        pass
    # GS image - teraz nas nie intersuje
    else:
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    # Obraz kolowowy
    if row['Fragments'] is not None:
        document.add_heading(f'Fragments {row['Fragments']}',3)
        # mamy nie pustą listę fragmentów
        for f in row['Fragments'][0]:

            fragment = img[f[0]:f[2],f[1]:f[3]].copy()
            # tu wykonujesz operacje i inne wyświetlenia na fragmencie
            memfile = plotImgFromCV(fragment)
            document.add_paragraph(f'fragment {f}')
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
            memfile.close()

document.save('report.docx')