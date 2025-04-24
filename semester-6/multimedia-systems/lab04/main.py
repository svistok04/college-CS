import time
import numpy as np
import cv2
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from glob import glob
import os


def colorFit(pixel, palette):
    idx = np.argmin(np.linalg.norm(palette - pixel, axis=1))
    return palette[idx].astype(np.float32)

def quantum_colorFit(img, palette):
    out_img = img.copy()
    for w in range(img.shape[0]):
        for k in range(img.shape[1]):
            out_img[w, k] = colorFit(img[w, k], palette)
    return out_img

def dithering_random(img):
    r = np.random.rand(*img.shape)
    return (img >= r).astype(np.float32)

def dithering_order(img, M, palette):
    n = int(M.shape[0] / 2)
    Mpre = (M + 1.0) / (2*n)**2 - 0.5

    out_img = img.copy()
    # r = 255 / len(palette)
    r = 1
    for x in range(img.shape[0]):
        for y in range(img.shape[1]):
            threshold = Mpre[x % (2 * n), y % (2 * n)]
            shifted_pixel = np.clip(img[x, y] + r * threshold, 0.0, 1.0)
            out_img[x, y] = colorFit(shifted_pixel, palette)

    return out_img


def dithering_fs(img, palette):
    height, width = img.shape[:2]
    out_img = img.copy()

    for y in range(height):
        for x in range(width):
            oldpixel = out_img[y, x]
            newpixel = colorFit(oldpixel, palette)
            out_img[y, x] = newpixel
            error = oldpixel - newpixel

            for dy, dx, coef in [(0, 1, 7 / 16), (1, -1, 3 / 16), (1, 0, 5 / 16), (1, 1, 1 / 16)]:
                ny, nx = y + dy, x + dx
                if 0 <= ny < height and 0 <= nx < width:
                    out_img[ny, nx] += error * coef

    return np.clip(out_img, 0.0, 1.0)


def make_palette(n_bits):
    levels = 2 ** n_bits
    return np.linspace(0, 1, levels).reshape(levels, 1)

def run_one_quantized_image(img, image_type, palette_name, palette_data):
    memfile = BytesIO()

    if image_type == "Grayscale":
        fig, axs = plt.subplots(2, 3, figsize=(12, 8))
        fig.suptitle(f"Dithering {palette_name}", fontsize=16)
        axs = axs.flatten()

        axs[0].imshow(img, cmap='gray', vmin=0, vmax=1)
        axs[0].set_title("Original")
        axs[0].axis('off')

        axs[1].imshow(quantum_colorFit(img, palette_data), cmap='gray', vmin=0, vmax=1)
        axs[1].set_title("Quantized")
        axs[1].axis('off')

        axs[2].imshow(dithering_funcs['Ordered'](img, M2, palette_data), cmap='gray', vmin=0, vmax=1)
        axs[2].set_title("Ordered")
        axs[2].axis('off')

        if palette_name == "1-bit":
            axs[3].imshow(dithering_funcs['Random'](img), cmap='gray', vmin=0, vmax=1)
            axs[3].set_title("Random")
        else:
            axs[3].axis('off')
        axs[3].axis('off') if axs[3].get_title() == '' else None

        axs[4].imshow(dithering_funcs['Floyd'](img, palette_data), cmap='gray', vmin=0, vmax=1)
        axs[4].set_title("Floyd")
        axs[4].axis('off')

        axs[5].axis('off')

    else:
        fig, axs = plt.subplots(2, 2, figsize=(10, 8))
        fig.suptitle(f"Dithering {palette_name}", fontsize=16)
        axs = axs.flatten()

        axs[0].imshow(img)
        axs[0].set_title("Original")
        axs[0].axis('off')

        axs[1].imshow(dithering_funcs['Ordered'](img, M2, palette_data))
        axs[1].set_title("Ordered")
        axs[1].axis('off')

        axs[2].imshow(quantum_colorFit(img, palette_data))
        axs[2].set_title("Quantized")
        axs[2].axis('off')

        axs[3].imshow(dithering_funcs['Floyd'](img, palette_data))
        axs[3].set_title("Floyd")
        axs[3].axis('off')

    fig.tight_layout()
    fig.savefig(memfile, format='png')
    plt.close(fig)
    memfile.seek(0)
    return memfile


palette8 = np.array([
    [0.0, 0.0, 0.0],
    [0.0, 0.0, 1.0],
    [0.0, 1.0, 0.0],
    [0.0, 1.0, 1.0],
    [1.0, 0.0, 0.0],
    [1.0, 0.0, 1.0],
    [1.0, 1.0, 0.0],
    [1.0, 1.0, 1.0],
])

palette16 =  np.array([
    [0.0, 0.0, 0.0,],
    [0.0, 1.0, 1.0,],
    [0.0, 0.0, 1.0,],
    [1.0, 0.0, 1.0,],
    [0.0, 0.5, 0.0,],
    [0.5, 0.5, 0.5,],
    [0.0, 1.0, 0.0,],
    [0.5, 0.0, 0.0,],
    [0.0, 0.0, 0.5,],
    [0.5, 0.5, 0.0,],
    [0.5, 0.0, 0.5,],
    [1.0, 0.0, 0.0,],
    [0.75, 0.75, 0.75,],
    [0.0, 0.5, 0.5,],
    [1.0, 1.0, 1.0,],
    [1.0, 1.0, 0.0,]
])

M2 = np.array([
    [0,  8,  2, 10],
    [12, 4, 14, 6],
    [3, 11, 1,  9],
    [15, 7, 13, 5]
])

dithering_funcs = {
    'Random': dithering_random,
    'Ordered': dithering_order,
    'Floyd': dithering_fs,
}
gray_palettes = {
    '1-bit': make_palette(1),
    '2-bit': make_palette(2),
    '4-bit': make_palette(4),
}

color_palettes = {
    'palette8': palette8,
    'palette16': palette16,
}



def main():


    df = pd.DataFrame([
        {
            'Filename': 'GS_0001.tif',
            'ImageType': 'Grayscale',
        },
        {
            'Filename': 'GS_0002.png',
            'ImageType': 'Grayscale',
        },
        {
            'Filename': 'GS_0003.png',
            'ImageType': 'Grayscale',
        },

        {
            'Filename': 'SMALL_0003.png',
            'ImageType': 'Color',
        },
        {
            'Filename': 'SMALL_0005.jpg',
            'ImageType': 'Color',
        },
        {
            'Filename': 'SMALL_0007.jpg',
            'ImageType': 'Color',
        },
        {
            'Filename': 'SMALL_0009.jpg',
            'ImageType': 'Color',
        },
    ])

    document = Document()
    document.add_heading('Quantization and Dithering Report', level=0)
    for idx, row in df.iterrows():
        filename = row['Filename']
        image_type = row['ImageType']
        print(filename)

        folder = '../IMG_GS' if image_type == 'Grayscale' else '../IMG_SMALL'
        path = os.path.join(folder, filename)

        img = cv2.imread(path, cv2.IMREAD_GRAYSCALE if image_type == 'Grayscale' else cv2.IMREAD_COLOR)
        img = img.astype(np.float32) / 255.0
        if image_type == 'Color':
            img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        document.add_page_break()
        document.add_heading(f"File: {filename}", level=1)

        if image_type == "Grayscale":
            for label, palette in gray_palettes.items():
                memfile = run_one_quantized_image(img, "Grayscale", label, palette)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()
        else:
            for label, palette in color_palettes.items():
                memfile = run_one_quantized_image(img, "Color", label, palette)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()

    document.save("quantization_dithering_report2.docx")

main()

#
# print(colorFit(0.43, palette_gray))
# print(colorFit(0.66, palette_gray))
# print(colorFit(0.8,  palette_gray))
#
# print(colorFit(np.array([0.25, 0.25, 0.5]), pallet8))
#
# print(colorFit(np.array([0.25, 0.25, 0.5]), pallet16))
#
# img = cv2.imread('../IMG_GS/GS_0001.tif', cv2.IMREAD_GRAYSCALE)
# img = img.astype(np.float32) / 255.0
# plt.imshow(img, cmap='gray')
# plt.show()
# palettes = {
#     1: make_palette(1),
#     2: make_palette(2),
#     4: make_palette(4),
# }
#
# quantized_imgs = {
#     bits: quantum_colorFit(img, palette) for bits, palette in palettes.items()
# }
#
# fig, axs = plt.subplots(1, 4, figsize=(16, 4))
# axs[0].imshow(img, cmap='gray', vmin=0, vmax=1)
# axs[0].set_title("Original")
# axs[0].axis('off')
#
# for i, (bits, img) in enumerate(quantized_imgs.items()):
#     axs[i+1].imshow(img, cmap='gray', vmin=0, vmax=1)
#     axs[i+1].set_title(f"{bits}-bit Quantization")
#     axs[i+1].axis('off')

# plt.show()
#
# img = cv2.imread('../IMG_GS/GS_0001.tif', cv2.IMREAD_GRAYSCALE)
# img = img.astype(np.float32) / 255.0
#
# dithered = dithering_random(img)
#
# plt.imshow(dithered, cmap='gray', vmin=0, vmax=1)
# plt.title("Random Dithering (1-bit)")
# plt.axis('off')
# plt.show()

# palette = np.linspace(0, 1, 2).reshape(-1, 1)
#
# M1 = np.array([
#     [0, 2],
#     [3, 1]
# ])
#
# M2 = np.array([
#     [0,  8,  2, 10],
#     [12, 4, 14, 6],
#     [3, 11, 1,  9],
#     [15, 7, 13, 5]
# ])
#
# M4 = np.block([
#     [4*M2 + 0, 4*M2 + 2],
#     [4*M2 + 3, 4*M2 + 1]
# ])

# img = cv2.imread('../IMG_GS/GS_0001.tif', cv2.IMREAD_GRAYSCALE)
# img = img.astype(np.float32) / 255.0
#
# ordered = dithering_order(img, M1, palette)
#
# plt.imshow(ordered, cmap='gray', vmin=0, vmax=1)
# plt.title("Ordered Dithering (1-bit)")
# plt.axis('off')
# plt.show()
# print(np.unique(ordered))  # Should show [0.0, 1.0]
#
# print('hi')
# print(np.unique(dithering_fs(img.copy(),np.linspace(0,1,2).reshape(2,1))).size) # ==2
# fs = dithering_fs(img, palette)
# plt.imshow(fs, cmap='gray')
# plt.title('dithering_fs')
# plt.show()
