import time
import numpy as np
import cv2
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from glob import glob
import os


def upscaling_nearest_neighbor(image, scaling_factor):

    if scaling_factor <= 1:
        print('Cannot upscale')
        return image
    height, width = image.shape[:2]
    new_height = np.ceil(height * scaling_factor).astype(int)
    new_width = np.ceil(width * scaling_factor).astype(int)
    X, Y = np.mgrid[0:height:new_height * 1j, 0:width:new_width * 1j]
    X = np.clip(np.floor(X).astype(int), 0, height - 1)
    Y = np.clip(np.floor(Y).astype(int), 0, width - 1)
    new_image = np.zeros((new_height, new_width, image.shape[2]), dtype=np.uint8)
    if image.shape[2] < 3:
        # grayscale
        pass
    else:
        # rgb
        new_image = image[X, Y]

    return new_image

def upscaling_bilinear_interpolation(image, scaling_factor):
    if scaling_factor <= 1:
        print('Cannot upscale')
        return image
    height, width = image.shape[:2]
    new_height = np.ceil(height * scaling_factor).astype(int)
    new_width = np.ceil(width * scaling_factor).astype(int)
    X, Y = np.mgrid[0:height:new_height * 1j, 0:width:new_width * 1j]
    X = np.floor(X).astype(int)
    Y = np.floor(Y).astype(int)

    row1 = np.clip(X + 1, 0, height - 1)
    col1 = np.clip(Y + 1, 0, width - 1)
    row0 = np.clip(X, 0, height - 1)
    col0 = np.clip(Y, 0, width - 1)

    row_weight = X - row0
    col_weight = Y - col0

    top_left = image[row0, col0]
    top_right = image[row0, col1]
    bottom_left = image[row1, col0]
    bottom_right = image[row1, col1]

    top = top_left * (1 - col_weight)[..., None] + top_right * col_weight[..., None]
    bottom = bottom_left * (1 - col_weight)[..., None] + bottom_right * col_weight[..., None]
    new_image = top * (1 - row_weight)[..., None] + bottom * row_weight[..., None]


    return new_image


def downscaling_mean_np(image, scaling_factor):
    if scaling_factor >= 1:
        print('Cannot downscale')
        return image

    height, width, channels = image.shape
    new_height = np.ceil(height * scaling_factor).astype(int)
    new_width = np.ceil(width * scaling_factor).astype(int)

    X, Y = np.mgrid[0:height, 0:width]
    mapped_X = (X * scaling_factor).astype(int)
    mapped_Y = (Y * scaling_factor).astype(int)

    new_image = np.zeros((new_height, new_width, channels), dtype=np.uint32)
    count = np.zeros((new_height, new_width), dtype=np.uint16)

    mapped_X_flat = mapped_X.ravel()
    mapped_Y_flat = mapped_Y.ravel()
    pixels_flat = image.reshape(-1, channels)

    np.add.at(new_image, (mapped_X_flat, mapped_Y_flat, slice(None)), pixels_flat)

    np.add.at(count, (mapped_X_flat, mapped_Y_flat), 1)

    count = np.maximum(count, 1)
    return (new_image / count[..., None]).astype(np.uint8)

def downscaling_mean(image, scaling_factor):
    height, width, channels = image.shape
    new_height = int(np.ceil(height * scaling_factor))
    new_width = int(np.ceil(width * scaling_factor))

    xx = np.linspace(0, height - 1, new_height)
    yy = np.linspace(0, width - 1, new_width)

    new_image = np.zeros((new_height, new_width, channels), dtype=np.uint8)

    for i in range(new_height):
        for j in range(new_width):
            if i > 0:
                x1 = -(xx[i] - xx[i - 1]) / 2
            else:
                x1 = 0
            if i < new_height - 1:
                x2 = (xx[i + 1] - xx[i]) / 2 + 1
            else:
                x2 = 0
            x_range = np.round(xx[i] + np.arange(x1, x2)).astype(int)
            x_range = np.clip(x_range, 0, height - 1)

            if j > 0:
                y1 = -(yy[j] - yy[j - 1]) / 2
            else:
                y1 = 0
            if j < new_width - 1:
                y2 = (yy[j + 1] - yy[j]) / 2 + 1
            else:
                y2 = 0
            y_range = np.round(yy[j] + np.arange(y1, y2)).astype(int)
            y_range = np.clip(y_range, 0, width - 1)

            fragment = image[np.ix_(x_range, y_range)]

            mean_pixel = np.mean(fragment, axis=(0, 1))
            new_image[i, j] = mean_pixel.astype(np.uint8)

    return new_image

def downscaling_median(image, scaling_factor):
    if scaling_factor >= 1:
        print('Cannot downscale')
        return image

    height, width, channels = image.shape
    new_height = int(np.ceil(height * scaling_factor))
    new_width = int(np.ceil(width * scaling_factor))

    xx = np.linspace(0, height - 1, new_height)
    yy = np.linspace(0, width - 1, new_width)

    new_image = np.zeros((new_height, new_width, channels), dtype=np.uint8)

    for i in range(new_height):
        for j in range(new_width):
            if i > 0:
                x1 = -(xx[i] - xx[i - 1]) / 2
            else:
                x1 = 0
            if i < new_height - 1:
                x2 = (xx[i + 1] - xx[i]) / 2 + 1
            else:
                x2 = 0
            x_range = np.round(xx[i] + np.arange(x1, x2)).astype(int)
            x_range = np.clip(x_range, 0, height - 1)

            if j > 0:
                y1 = -(yy[j] - yy[j - 1]) / 2
            else:
                y1 = 0
            if j < new_width - 1:
                y2 = (yy[j + 1] - yy[j]) / 2 + 1
            else:
                y2 = 0
            y_range = np.round(yy[j] + np.arange(y1, y2)).astype(int)
            y_range = np.clip(y_range, 0, width - 1)

            fragment = image[np.ix_(x_range, y_range)]

            median_pixel = np.median(fragment, axis=(0, 1))
            new_image[i, j] = median_pixel.astype(np.uint8)

    return new_image

def downscaling_weighted_mean(image, scaling_factor):
    height, width, channels = image.shape
    new_height = int(np.ceil(height * scaling_factor))
    new_width = int(np.ceil(width * scaling_factor))

    xx = np.linspace(0, height - 1, new_height)
    yy = np.linspace(0, width - 1, new_width)

    new_image = np.zeros((new_height, new_width, channels), dtype=np.uint8)

    for i in range(new_height):
        for j in range(new_width):
            if i > 0:
                x1 = -(xx[i] - xx[i - 1]) / 2
            else:
                x1 = 0
            if i < new_height - 1:
                x2 = (xx[i + 1] - xx[i]) / 2 + 1
            else:
                x2 = 0
            x_range = np.round(xx[i] + np.arange(x1, x2)).astype(int)
            x_range = np.clip(x_range, 0, height - 1)

            if j > 0:
                y1 = -(yy[j] - yy[j - 1]) / 2
            else:
                y1 = 0
            if j < new_width - 1:
                y2 = (yy[j + 1] - yy[j]) / 2 + 1
            else:
                y2 = 0
            y_range = np.round(yy[j] + np.arange(y1, y2)).astype(int)
            y_range = np.clip(y_range, 0, width - 1)

            fragment = image[np.ix_(x_range, y_range)]

            weights = np.random.rand(*fragment.shape[:2])
            weights_sum = np.sum(weights)
            if weights_sum == 0:
                weights = np.ones_like(weights)
                weights_sum = np.sum(weights)

            weighted = np.zeros((channels,), dtype=np.float64)
            for c in range(channels):
                weighted[c] = np.sum(fragment[:, :, c] * weights)

            weighted_pixel = weighted / weights_sum
            new_image[i, j] = weighted_pixel.astype(np.uint8)

    return new_image

def canny_edges(img, low_threshold=50, high_threshold=150):
    if img.dtype != np.uint8:
        if img.max() <= 1.0:
            img = (img * 255).astype(np.uint8)
        else:
            img = img.astype(np.uint8)

    if img.ndim == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    else:
        gray = img

    edges = cv2.Canny(gray, low_threshold, high_threshold)
    return edges



def get_fragment(img, coords):
    x1, y1, x2, y2 = coords
    return img[y1:y2, x1:x2]

def run_one_image(img, algos: dict, scaling_factor: float):
    results = {}
    results["Original"] = img
    results["Canny(Original)"] = canny_edges(img)

    times = {}
    for name, algo in algos.items():
        t = time.time()
        out = algo(img, scaling_factor)
        times[name] = time.time() - t
        results[name] = out
        results[f"Canny({name})"] = canny_edges(out)

    top_row_imgs = [results[k] for k in results if not k.startswith("Canny")]
    top_row_titles = [k for k in results if not k.startswith("Canny")]

    bottom_row_imgs = [results[k] for k in results if k.startswith("Canny")]
    bottom_row_titles = [k for k in results if k.startswith("Canny")]

    fig, axs = plt.subplots(2, len(top_row_imgs), figsize=(5 * len(top_row_imgs), 10))
    for i in range(len(top_row_imgs)):
        axs[0, i].imshow(top_row_imgs[i])
        axs[0, i].set_title(top_row_titles[i])
        axs[0, i].axis('off')

        axs[1, i].imshow(bottom_row_imgs[i], cmap='gray')
        axs[1, i].set_title(bottom_row_titles[i])
        axs[1, i].axis('off')

    memfile = BytesIO()
    fig.tight_layout()
    fig.savefig(memfile, format='png')
    plt.close(fig)
    return memfile, times


# def get_all_images(folder):
#     exts = ['*.jpg', '*.jpeg', '*.png', '*.tif', '*.tiff']
#     files = []
#     for ext in exts:
#         files.extend(glob(f"{folder}/{ext}"))
#     return sorted(files)

df = pd.DataFrame([
    {
        'Filename': 'BIG_0001.jpg',
        'Fragments': [[500, 2000, 900, 2400],
                      [1000, 1000, 1500, 1600],
                      [3500, 600, 4000, 1200]],
        'Type': 'Downscaling',
        'ScaleFactors': [0.1, 0.05],
    },
    {
        'Filename': 'BIG_0002.jpg',
        'Fragments': [[1000, 2400, 1600, 3000],
                      [800, 1200, 1400, 1800],
                      [3000, 1800, 3600, 2400]],
        'Type': 'Downscaling',
        'ScaleFactors': [0.1, 0.05],
    },
    {
        'Filename': 'SMALL_0003.png',
        'Fragments': None,
        'Type': 'Upscaling',
        'ScaleFactors': [2.0, 5.0],
    },
    {
        'Filename': 'SMALL_0005.jpg',
        'Fragments': None,
        'Type': 'Upscaling',
        'ScaleFactors': [2.0, 5.0],
    },
    {
        'Filename': 'SMALL_0007.jpg',
        'Fragments': None,
        'Type': 'Upscaling',
        'ScaleFactors': [2.0, 5.0],
    },
    {
        'Filename': 'SMALL_0009.jpg',
        'Fragments': None,
        'Type': 'Upscaling',
        'ScaleFactors': [2.0, 5.0],
    }
])

def main():
    # files_downscaling = get_all_images('IMG_BIG')
    # files_upscaling = get_all_images('IMG_SMALL')



    downscaling_algos = {
        'Mean': downscaling_mean,
        'Median': downscaling_median,
        'WeightedMean': downscaling_weighted_mean,
    }

    upscaling_algos = {
        'Nearest neighbors': upscaling_nearest_neighbor,
        'Bilinear interpolation': upscaling_bilinear_interpolation,
    }

    document = Document()
    document.add_heading('Zmień ten tytuł', level=0)

    document.add_heading('Downscaling Results', level=1)

    for idx, row in df.iterrows():
        filename = row['Filename']
        fragments = row['Fragments']
        img_type = row['Type']
        scales = row['ScaleFactors']

        folder = 'IMG_BIG' if img_type == 'Downscaling' else 'IMG_SMALL'
        path = os.path.join(folder, filename)
        img = plt.imread(path)

        document.add_heading(f"File: {filename}", level=2)

        if fragments:
            document.add_heading('Downscaling Results', level=1)
            for i, fragment_coords in enumerate(fragments):
                fragment = get_fragment(img, fragment_coords)
                document.add_heading(f"Fragment {i+1} – Coords: {fragment_coords}", level=2)

                for scale in scales:
                    document.add_heading(f"Scale: {int(scale * 100)}%", level=3)
                    memfile, times = run_one_image(fragment, downscaling_algos, scale)
                    document.add_picture(memfile, width=Inches(6))
                    memfile.close()

                    for algo, t in times.items():
                        document.add_paragraph(f"{algo}: {t:.2f} seconds")
        else:
            document.add_page_break()
            document.add_heading('Upscaling Results', level=1)
            for scale in scales:
                document.add_heading(f"Scale: {int(scale * 100)}%", level=2)
                memfile, times = run_one_image(img, upscaling_algos, scale)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()

                for algo, t in times.items():
                    document.add_paragraph(f"{algo}: {t:.2f} seconds")

    document.save("image_scaling_report.docx")

main()

# img = plt.imread("IMG_SMALL/SMALL_0003.png")
# plt.imshow(img)
# plt.show()
#
# img = upscaling_bilinear_interpolation(img, 1.5)
# plt.imshow(img)
# plt.show()



#
# img = plt.imread('IMG_SMALL/SMALL_0009.jpg')
# # img= np.zeros((11,11,3),dtype=np.uint8)
# # img[5,5,:] = 255
# fig = plt.figure(figsize=(16, 4))
# ax_original = fig.add_subplot(1, 3, 1)
# ax_original.imshow(img)
# # plt.imshow(img)
# # plt.show()
# # print(np.max(img))
# # height, width = img.shape[:2]
# height, width = 2, 2
# size = height * width
# scaling_factor = 2.0
# new_height = int(height * scaling_factor)
# new_width = int(width * scaling_factor)
# new_size = new_height * new_width
# print(new_height, new_width)
# X, Y = np.mgrid[0:height:new_height * 1j, 0:width:new_width * 1j]
# print(X)
# print(Y)
# f = time.time()
# img_nn = upscaling_nearest_neighbor(img, 5.0)
# # print(img.shape)
# ax_nn = fig.add_subplot(1, 3, 2)
# ax_nn.imshow(img_nn)
# ax_bi = fig.add_subplot(1, 3, 3)
# img_bi = upscaling_bilinear_interpolation(img, 5.0)
# ax_bi.imshow(img_bi)
# plt.show()
# print(time.time() - f)
#
# img = plt.imread('IMG_BIG/BIG_0002.jpg')
# # t = time.time()
# plt.imshow(img)
# plt.show()
# # img_mean = downscaling_mean_np(img, 0.03)
# # plt.imshow(img_mean)
# # plt.show()
# # print(f'it took {time.time()-t} seconds for mine')
#
# t = time.time()
# img_mean = downscaling_mean(img, 0.03)
# plt.imshow(img_mean)
# plt.show()
# print(f'it took {time.time()-t} seconds for for nested loops')
#
# t = time.time()
# img_median = downscaling_median(img, 0.03)
# plt.imshow(img_median)
# plt.show()
# print(f'it took {time.time()-t} seconds for nested loops')
#
# t = time.time()
# img_weighted_mean = downscaling_weighted_mean(img, 0.03)
# plt.imshow(img_weighted_mean)
# plt.show()