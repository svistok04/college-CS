import cv2
import numpy as np
import sys
from docx import Document
from docx.shared import T, Inches
from io import BytesIO
import matplotlib.pyplot as plt
import scipy.fftpack
import os
import time

class ver2:
    def __init__(self, Y, Cb, Cr, OGShape, Y_shape, Cb_shape, Cr_shape, Ratio="4:4:4", QY=np.ones((8,8)), QC=np.ones((8,8))):
        self.shape = OGShape
        self.Y_shape = Y_shape
        self.Cb_shape = Cb_shape
        self.Cr_shape = Cr_shape
        self.Y = Y
        self.Cb = Cb
        self.Cr = Cr
        self.ChromaRatio = Ratio
        self.QY = QY
        self.QC = QC

def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

def compressJPG(img, ratio="4:4:4", QY=np.ones((8, 8)), QC=np.ones((8, 8))):
    img = cv2.cvtColor(img, cv2.COLOR_RGB2YCrCb).astype(int)
    OGShape = img.shape

    Y, Cr, Cb = chroma_subsampling(img, ratio)

    Y_comp = compressLayer(Y, QY)
    Cr_comp = compressLayer(Cr, QC)
    Cb_comp = compressLayer(Cb, QC)

    return ver2(Y_comp, Cb_comp, Cr_comp, OGShape, Y.shape, Cb.shape, Cr.shape, Ratio=ratio, QY=QY, QC=QC)

def compressLayer(layer, Q):
    S = np.array([])
    for w in range(0, layer.shape[0], 8):
        for k in range(0, layer.shape[1], 8):
            block = layer[w:(w+8), k:(k+8)]
            S = np.append(S, compressBlock(block, Q))
    return byterun_encode(S)


def compressBlock(block, Q):
    block -= 128
    dct_block = dct2(block)

    quantized = np.round(dct_block / Q).astype(np.int8)

    zigzagged = zigzag(quantized)
    
    return zigzagged


def decompressJPG(JPEG):
    Y = decompressLayer(JPEG.Y, JPEG.QC, JPEG.Y_shape)
    Cr = decompressLayer(JPEG.Cr, JPEG.QC, JPEG.Cr_shape)
    Cb = decompressLayer(JPEG.Cb, JPEG.QC, JPEG.Cb_shape)

    Cr, Cb = chroma_resampling(Cr, Cb, JPEG.ChromaRatio)

    img = np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)
    img = cv2.cvtColor(img, cv2.COLOR_YCrCb2RGB)
    return img

def decompressBlock(vector, Q):
    matrix = zigzag(vector)
    
    dequantized = matrix * Q
    
    block = idct2(dequantized)
    
    return block + 128

def decompressLayer(S, Q, original_shape):
    decoded = byterun_decode(S)

    h, w = original_shape

    L = np.zeros((h, w))

    idx = 0
    for y in range(0, h, 8):
        for x in range(0, w, 8):
            vector = decoded[idx * 64 : (idx + 1) * 64]
            L[y:y+8, x:x+8] = decompressBlock(vector, Q)
            idx += 1

    return L

def dct2(a):
    return scipy.fftpack.dct(scipy.fftpack.dct(a.astype(float), axis=0, norm='ortho'), axis=1, norm='ortho')

def idct2(a):
    return scipy.fftpack.idct(scipy.fftpack.idct(a.astype(float), axis=0, norm='ortho'), axis=1, norm='ortho')



def zigzag(A):
    template = np.array([
        [0,  1,  5,  6,  14, 15, 27, 28],
        [2,  4,  7,  13, 16, 26, 29, 42],
        [3,  8,  12, 17, 25, 30, 41, 43],
        [9,  11, 18, 24, 31, 40, 44, 53],
        [10, 19, 23, 32, 39, 45, 52, 54],
        [20, 22, 33, 38, 46, 51, 55, 60],
        [21, 34, 37, 47, 50, 56, 59, 61],
        [35, 36, 48, 49, 57, 58, 62, 63],
    ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B



def chroma_subsampling(img, ratio="4:4:4"):
    Y = img[:, :, 0]
    Cr = img[:, :, 1]
    Cb = img[:, :, 2]

    if ratio == "4:2:2":
        Cr = Cr[:, ::2]
        Cb = Cb[:, ::2]
    elif ratio == "4:4:4":
        pass 

    return Y, Cr, Cb

def chroma_resampling(Cr, Cb, ratio="4:4:4"):
    if ratio == "4:2:2":
        Cr = np.repeat(Cr, 2, axis=1)
        Cb = np.repeat(Cb, 2, axis=1)
    elif ratio == "4:4:4":
        pass

    return Cr, Cb

def byterun_encode(data):
    shape = data.shape
    if len(shape) > 1:
        data = data.flatten()

    length = len(data)
    result = np.zeros(length * 2, dtype=np.uint8)
    out_idx = 0
    i = 0

    while i < length:
        run_start = i
        run_value = data[i]
        run_length = 1
        while i + run_length < length and data[i + run_length] == run_value and run_length < 127:
            # print(run_value)
            run_length += 1

        if run_length > 1:
            # print(f'repeated run length {run_length}')
            result[out_idx] = (-run_length + 256) % 256
            result[out_idx + 1] = run_value
            out_idx += 2
            i += run_length
        else:
            literal_start = i
            literal_length = 1
            while i + literal_length < length - 1 and (data[i + literal_length] != data[i + literal_length + 1] or literal_length == 1) and literal_length < 127:
                literal_length += 1
            result[out_idx] = literal_length
            result[out_idx + 1:out_idx + literal_length + 1] = data[i:i + literal_length]
            out_idx += 1 + literal_length
            i += literal_length

    return np.concatenate((
        np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.uint8),
                result[:out_idx].view(np.uint8)
    ))
    # return np.concatenate((
    #     np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.int8),
    #     result[:out_idx]
    # ))

def byterun_decode(encoded):
    ndim = encoded[:4].view(np.uint32)[0]
    shape = encoded[4:4 + 4 * ndim].view(np.uint32)
    data_start = 4 + 4 * ndim

    data = encoded[data_start:].view(np.int8)
    result = []

    i = 0
    while i < len(data):
        count = data[i]
        if count < 0:
            value = data[i + 1]
            result.extend([value] * (-count))
            i += 2
        else:
            count = int(count)
            result.extend(data[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(result, dtype=np.int8).reshape(tuple(shape))

QY = np.array([
    [16,11,10,16,24,40,51,61],
    [12,12,14,19,26,58,60,55],
    [14,13,16,24,40,57,69,56],
    [14,17,22,29,51,87,80,62],
    [18,22,37,56,68,109,103,77],
    [24,36,55,64,81,104,113,92],
    [49,64,78,87,103,121,120,101],
    [72,92,95,98,112,100,103,99]
],)
QC = np.array([
    [17,18,24,47,99,99,99,99],
    [18,21,26,66,99,99,99,99],
    [24,26,56,99,99,99,99,99],
    [47,66,99,99,99,99,99,99],
    [99,99,99,99,99,99,99,99],
    [99,99,99,99,99,99,99,99],
    [99,99,99,99,99,99,99,99],
    [99,99,99,99,99,99,99,99]
])

def plot_comparison(original, decompressed, title):
    fig, axs = plt.subplots(1, 2, figsize=(8, 4))
    axs[0].imshow(original)
    axs[0].set_title("Original")
    axs[0].axis("off")
    axs[1].imshow(decompressed)
    axs[1].set_title("Decompressed")
    axs[1].axis("off")
    plt.suptitle(title)
    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    return buf

def main():
    filenames = ['./color1.jpg', './document1.jpg', './drawing1.jpg']
    ratios = ["4:4:4", "4:2:2"]
    Qs = [("(QN, QN)", np.ones((8,8)), np.ones((8,8))),
          ("(QY, QC)", QY, QC)]
    patch_size = 512

    document = Document()
    document.add_heading("JPEG Compression Report", level=0)

    results = []
    figures = []

    for filename in filenames:
        img = cv2.imread(filename)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        h, w, _ = img.shape
        patch = img[h//4:h//4+patch_size, w//4:w//4+patch_size]

        for ratio in ratios:
            for qlabel, QY_used, QC_used in Qs:
                start = time.time()
                compressed = compressJPG(patch, ratio, QY_used, QC_used)
                t_comp = time.time() - start

                start = time.time()
                decompressed = decompressJPG(compressed)
                t_decomp = time.time() - start

                original_size = get_size(patch)
                compressed_size = get_size(compressed)
                compression_percent = (compressed_size / original_size) * 100

                fig_title = f"{os.path.basename(filename)} {ratio} - {qlabel}"
                fig_buf = plot_comparison(patch, decompressed, fig_title)
                figures.append((filename, fig_title, fig_buf))

                results.append((os.path.basename(filename), ratio, qlabel,
                                f"{compression_percent:.1f}%",
                                f"{t_comp*1000:.1f} ms",
                                f"{t_decomp*1000:.1f} ms"))

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image'
    hdr_cells[1].text = 'Ratio'
    hdr_cells[2].text = 'Q Table'
    hdr_cells[3].text = 'Size After (%)'
    hdr_cells[4].text = 'Compress Time'
    hdr_cells[5].text = 'Decompress Time'

    for row in results:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    document.add_page_break()
    last_file = None
    for filename, fig_title, fig_buf in figures:
        if filename != last_file:
            document.add_heading(f"File: {os.path.basename(filename)}", level=2)
            last_file = filename
        document.add_picture(fig_buf, width=Inches(5.5))

    document.save("jpeg_report.docx")

main()
