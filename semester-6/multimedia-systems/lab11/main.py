from io import BytesIO
from docx import Document
from docx.shared import T, Inches
import matplotlib.pyplot as plt
import cv2
import numpy as np
from skimage.metrics import peak_signal_noise_ratio as psnr, structural_similarity as ssim
import os

def water_mark(img, mask, alpha=0.25):
    assert img.shape[:2] == mask.shape[:2], "Wrong size"
    flag = len(img.shape) < 3
    t_img = cv2.cvtColor(img, cv2.COLOR_GRAY2RGBA) if flag else cv2.cvtColor(img, cv2.COLOR_RGB2RGBA)
    if mask.dtype == bool:
        t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
    elif mask.dtype == np.uint8:
        t_mask = cv2.cvtColor(mask.astype(np.uint8), cv2.COLOR_GRAY2RGBA)
    else:
        t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
    t_out = cv2.addWeighted(t_img, 1, t_mask, alpha, 0)
    out = cv2.cvtColor(t_out, cv2.COLOR_RGBA2GRAY) if flag else cv2.cvtColor(t_out, cv2.COLOR_RGBA2RGB)
    return out

def put_data(img, data, binary_mask=np.uint8(1)):
    assert img.dtype == np.uint8, "img wrong data type"
    assert binary_mask.dtype == np.uint8, "binary_mask wrong data type"
    un_binary_mask = np.unpackbits(np.array([binary_mask], dtype=np.uint8))
    unpacked_data = np.unpackbits(data) if data.dtype == bool else data
    dataspace = img.shape[0] * img.shape[1] * np.sum(un_binary_mask)
    assert dataspace == unpacked_data.size, "too much data"
    prepered_data = unpacked_data.reshape((img.shape[0], img.shape[1], np.sum(un_binary_mask)))
    mask = np.full((img.shape[0], img.shape[1]), binary_mask, dtype=np.uint8)
    img = np.bitwise_and(img, np.invert(mask))
    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            temp = prepered_data[:, :, bv]
            temp = np.left_shift(temp, i)
            img = np.bitwise_or(img, temp)
            bv += 1
    return img

def pop_data(img, binary_mask=np.uint8(1), out_shape=None):
    un_binary_mask = np.unpackbits(np.array([binary_mask], dtype=np.uint8))
    data = np.zeros((img.shape[0], img.shape[1], np.sum(un_binary_mask)), dtype=np.uint8)
    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            mask = np.full((img.shape[0], img.shape[1]), 2 ** i, dtype=np.uint8)
            temp = np.bitwise_and(img, mask) >> i
            data[:, :, bv] = temp[:, :]
            bv += 1
    flat = data.flatten()
    if out_shape is not None:
        packed = np.packbits(flat)
        packed = packed[:np.prod(out_shape)]
        data_out = packed.reshape(out_shape)
        return data_out
    else:
        return flat

def add_table(document, title, rows, header):
    document.add_heading(title, level=2)
    table = document.add_table(rows=1+len(rows), cols=len(header))
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(header):
        hdr_cells[i].text = col_name
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    document.add_paragraph('')

def scenario_1_hide_text(document):
    cover1 = cv2.cvtColor(cv2.imread('cover1.png'), cv2.COLOR_BGR2RGB)
    with open('quote.txt', 'r', encoding='utf-8') as f:
        text = f.read()
    text_bytes = np.frombuffer(text.encode('utf-8'), dtype=np.uint8)
    bits = np.unpackbits(text_bytes)
    B = cover1[:,:,2].copy()
    dataspace = B.size
    if bits.size < dataspace:
        bits = np.pad(bits, (0, dataspace - bits.size), 'constant')
    elif bits.size > dataspace:
        bits = bits[:dataspace]
    B_encoded = put_data(B.copy(), bits, np.uint8(1))
    rec_bits = pop_data(B_encoded, np.uint8(1))
    rec_bytes = np.packbits(rec_bits[:len(text_bytes)*8])
    rec_text = rec_bytes.tobytes().decode('utf-8', errors='ignore')
    print(f'Original text: {text}')
    print(f'Recovered text: {rec_text}')
    print(f'Match? {text == rec_text}')
    psnr_val = psnr(B, B_encoded, data_range=255)
    ssim_val = ssim(B, B_encoded, data_range=255)

    _, axs = plt.subplots(1, 2, figsize=(8, 3))
    axs[0].imshow(B, cmap='gray')
    axs[0].set_title('Blue Channel')
    axs[0].axis('off')
    axs[1].imshow(B_encoded, cmap='gray')
    axs[1].set_title('Blue w/ Text')
    axs[1].axis('off')
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150)
    plt.close()
    buf.seek(0)

    document.add_heading("Steganography Hidden Text", level=2)
    document.add_picture(buf, width=Inches(5))
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Channel"
    table.rows[0].cells[1].text = "PSNR"
    table.rows[0].cells[2].text = "SSIM"
    table.rows[1].cells[0].text = "Blue"
    table.rows[1].cells[1].text = f"{psnr_val:.2f}"
    table.rows[1].cells[2].text = f"{ssim_val:.4f}"
    document.add_page_break()

def scenario_2_hide_image(document):
    cover2 = cv2.cvtColor(cv2.imread('cover3.png'), cv2.COLOR_BGR2RGB)
    cover3 = cv2.cvtColor(cv2.imread('cover2.png'), cv2.COLOR_BGR2RGB)
    bits_per_channel = [2, 2, 3]
    masks = [np.uint8((1 << bits_per_channel[0]) - 1), np.uint8((1 << bits_per_channel[1]) - 1), np.uint8((1 << bits_per_channel[2]) - 1)]
    encoded = cover2.copy()

    quantized_secret = np.zeros_like(cover3)
    for c, bits_pp in enumerate(bits_per_channel):
        levels = (1 << bits_pp) - 1
        quantized_secret[:,:,c] = np.round(cover3[:,:,c] / 255 * levels).astype(np.uint8)

    for c, (mask, bits_pp) in enumerate(zip(masks, bits_per_channel)):
        vals = quantized_secret[:,:,c]
        bits = np.unpackbits(vals[..., None], axis=2, bitorder='big')[:, :, -bits_pp:]
        print(np.shape(bits))
        bits_flat = bits.reshape(-1)
        encoded[:,:,c] = put_data(cover2[:,:,c], bits_flat, mask)

    recovered = np.zeros_like(cover3)
    for c, (mask, bits_pp) in enumerate(zip(masks, bits_per_channel)):
        flat_bits = pop_data(encoded[:,:,c], mask)
        bits = flat_bits.reshape(cover3.shape[0], cover3.shape[1], bits_pp)
        vals = np.zeros((cover3.shape[0], cover3.shape[1]), dtype=np.uint8)
        for b in range(bits_pp):
            vals |= (bits[:,:,b] << (bits_pp - b - 1))
        levels = (1 << bits_pp) - 1
        recovered[:,:,c] = np.round(vals / levels * 255).astype(np.uint8)

    metrics = []
    for i, ch in enumerate(['R','G','B']):
        ps = psnr(cover2[:,:,i], encoded[:,:,i], data_range=255)
        ss = ssim(cover2[:,:,i], encoded[:,:,i], data_range=255)
        metrics.append([ch, f"{ps:.2f}", f"{ss:.4f}"])

    _, axs = plt.subplots(1, 4, figsize=(16,5))
    axs[0].imshow(cover2); axs[0].set_title("Carrier (cover3.png)")
    axs[1].imshow(encoded); axs[1].set_title("Encoded (cover3.png+cover2.png)")
    axs[2].imshow(cover3); axs[2].set_title("Hidden (cover2.png)")
    axs[3].imshow(recovered); axs[3].set_title("Recovered (decoded cover2.png)")
    for ax in axs: ax.axis('off')
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150)
    plt.close()
    buf.seek(0)

    document.add_heading("Steganography Hidden Image", level=2)
    document.add_picture(buf, width=Inches(6))
    table = document.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "Kanał"
    table.rows[0].cells[1].text = "PSNR"
    table.rows[0].cells[2].text = "SSIM"
    for i, row in enumerate(metrics):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = str(val)
    document.add_page_break()

def scenario_3_watermark(document):
    cover1 = cv2.cvtColor(cv2.imread('cover1.png'), cv2.COLOR_BGR2RGB)
    mask = cv2.imread('logo.png', 0)
    alpha_vals = [0.10, 0.25, 0.5]
    images = []
    metrics = []
    for alpha in alpha_vals:
        watermarked = water_mark(cover1, mask, alpha)
        images.append(watermarked)
        ps = psnr(cover1, watermarked, data_range=255)
        ss = ssim(cover1, watermarked, channel_axis=2, data_range=255)
        metrics.append([f"{alpha:.2f}", f"{ps:.2f}", f"{ss:.4f}"])

    _, axs = plt.subplots(1, 2+len(alpha_vals), figsize=(14,4))
    axs[0].imshow(cover1); axs[0].set_title("Cover")
    axs[1].imshow(mask, cmap='gray'); axs[1].set_title("Mask")
    for i, alpha in enumerate(alpha_vals):
        axs[2+i].imshow(images[i])
        axs[2+i].set_title(f"WM α={alpha}")
    for ax in axs: ax.axis('off')
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150)
    plt.close()
    buf.seek(0)
    document.add_heading("Watermark", level=2)
    document.add_picture(buf, width=Inches(6))
    table = document.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "Alpha"
    table.rows[0].cells[1].text = "PSNR"
    table.rows[0].cells[2].text = "SSIM"
    for i, row in enumerate(metrics):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = str(val)
    document.add_page_break()

def scenario_4_bit_budget(document):
    cover2 = cv2.cvtColor(cv2.imread('cover2.png'), cv2.COLOR_BGR2RGB)
    scenarios = [
        (1, 1, 1),
        (2, 2, 2),
        (3, 3, 2),
        (4, 3, 4),
        (5, 5, 5),
        (7, 7, 7),
        (8, 0, 0),
        (0, 8, 0),
        (0, 0, 8)
    ]
    rows = []
    images = [cover2]
    labels = ['Original']
    for r_bits, g_bits, b_bits in scenarios:
        masks = [np.uint8((1<<r_bits)-1), np.uint8((1<<g_bits)-1), np.uint8((1<<b_bits)-1)]
        bits_per_channel = [r_bits, g_bits, b_bits]
        encoded = cover2.copy()
        for c, (mask, bits_pp) in enumerate(zip(masks, bits_per_channel)):
            if bits_pp == 0:
                continue
            dataspace = cover2[:,:,c].size * bits_pp
            rand_data = np.random.randint(0, 2, size=dataspace, dtype=np.uint8)
            encoded[:,:,c] = put_data(cover2[:,:,c], rand_data, mask)
        images.append(encoded.copy())
        labels.append(f"{r_bits}{g_bits}{b_bits} bits (R/G/B)")
        ps = psnr(cover2, encoded, data_range=255)
        ss = ssim(cover2, encoded, channel_axis=2, data_range=255)
        rows.append([f"{r_bits}/{g_bits}/{b_bits}", f"{ps:.2f}", f"{ss:.4f}"])
    
    n = len(images)
    ncols = 5
    nrows = (n + ncols - 1) // ncols
    _, axs = plt.subplots(nrows, ncols, figsize=(3*ncols, 4*nrows))
    axs = axs.flatten()
    for i in range(n):
        axs[i].imshow(images[i])
        axs[i].set_title(labels[i])
        axs[i].axis('off')
    for i in range(n, len(axs)):
        axs[i].axis('off')
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150)
    plt.close()
    buf.seek(0)

    document.add_heading("Image Deconstruction", level=2)
    document.add_picture(buf, width=Inches(6))
    table = document.add_table(rows=len(rows)+1, cols=3)
    table.rows[0].cells[0].text = "Bits (R/G/B)"
    table.rows[0].cells[1].text = "PSNR"
    table.rows[0].cells[2].text = "SSIM"
    for r, row in enumerate(rows):
        for c in range(3):
            table.rows[r+1].cells[c].text = row[c]
    document.add_page_break()

document = Document()
document.add_heading("Steganography and Watermarks Report", 0)
document.add_page_break()

scenario_1_hide_text(document)
scenario_2_hide_image(document)
scenario_3_watermark(document)
scenario_4_bit_budget(document)

document.save("Steganography_and_Watermarks_Report.docx")
print("!")
