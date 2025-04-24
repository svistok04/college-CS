import time
import numpy as np
import cv2
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from glob import glob
import os
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
from scipy.interpolate import interp1d

def plot_audio(signal, fs, time_margin=[0, 0.02], fsize=2 ** 15):
    fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    fig.suptitle(f'Time margin {time_margin}')
    fig.tight_layout(pad=1.5)

    total_duration = len(signal) / fs
    t_margin = [min(time_margin[0], total_duration), min(time_margin[1], total_duration)]
    time = np.linspace(0, total_duration, len(signal))
    axs[0].plot(time, signal)
    axs[0].set_xlim(t_margin[0], t_margin[1])
    axs[0].set_xlabel('Time (s)')
    axs[0].set_ylabel('Amplitude')

    fsize = min(fsize, len(signal))
    if fsize < 2:
        spectrum = np.array([0])
        x = np.array([0])
    else:
        yf = scipy.fftpack.fft(signal, fsize)
        spectrum = 20 * np.log10(np.abs(yf[:fsize // 2]) + np.finfo(np.float32).eps)
        x = np.linspace(0, fs / 2, len(spectrum))

    axs[1].plot(x, spectrum)
    axs[1].set_xlabel('Frequency (Hz)')
    axs[1].set_ylabel('Magnitude [dB]')

    # Save to memory
    memfile = BytesIO()
    fig.savefig(memfile, format='png')
    plt.close(fig)  # Prevent interactive display flood

    # Return main frequency and amplitude
    max_idx = np.argmax(spectrum)
    max_freq = float(x[max_idx]) if len(x) else 0.0
    max_amplitude = float(spectrum[max_idx]) if len(spectrum) else 0.0

    return memfile, max_freq, max_amplitude


def quantization(data, bit):
    dataF = data.astype(np.float32)

    if np.issubdtype(data.dtype, np.floating):
        data_min = np.min(dataF)
        data_max = np.max(dataF)
    else:
        info = np.iinfo(data.dtype)
        data_min = info.min
        data_max = info.max

    if data_max == data_min:
        return data.copy()

    data_norm = (dataF - data_min) / (data_max - data_min)

    levels = 2 ** bit
    data_q = np.round(data_norm * (levels - 1)) / (levels - 1)

    data_rescaled = data_q * (data_max - data_min) + data_min

    return data_rescaled.astype(data.dtype)

def decimate(data, n, fs):
    return data[::n], fs // n


def interpolate(signal, fs_old, fs_new, method='linear'):
    t_old = np.linspace(0, len(signal)/fs_old, len(signal))
    N_new = int(len(signal) * fs_new / fs_old)
    t_new = np.linspace(0, len(signal)/fs_old, N_new)

    interp_func = interp1d(t_old, signal, kind=method)
    new_signal = interp_func(t_new).astype(signal.dtype)

    return new_signal


# data = np.linspace(-8, 1, 255)
# print(np.unique(quantization(data, 3)))

from collections import defaultdict

results = defaultdict(list)


def main():
    bit_depths = [4, 8, 16, 24]
    decimation_steps = [2, 4, 6, 10, 24]
    interp_fs = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
    files = [
        'SIN/sin_60Hz.wav',
        'SIN/sin_440Hz.wav',
        'SIN/sin_8000Hz.wav',
        'SIN/sin_combined.wav'
    ]

    time_margins = [[0, 0.02],
                    [0, 0.02],
                    [0, 0.02],
                    [0, 0.02],]

    results = defaultdict(list)

    for file_path, time_margin in zip(files, time_margins):
        filename = os.path.basename(file_path)
        data, fs = sf.read(file_path)
        entry_list = []

        memfile, freq, amp = plot_audio(data, fs, time_margin)
        entry_list.append(("Original", memfile, freq, amp))

        for bits in bit_depths:
            q_data = quantization(data, bits)
            memfile, freq, amp = plot_audio(q_data, fs, time_margin)
            entry_list.append((f"{bits}-bit quant", memfile, freq, amp))

        for step in decimation_steps:
            d_data, new_fs = decimate(data, step, fs)
            memfile, freq, amp = plot_audio(d_data, new_fs, time_margin)
            entry_list.append((f"Decimation {step}", memfile, freq, amp))

        for target_fs in interp_fs:
            for method in ['linear', 'cubic']:
                i_data = interpolate(data, fs, target_fs, method)
                memfile, freq, amp = plot_audio(i_data, target_fs, time_margin)
                entry_list.append((f"{method} interp {target_fs}Hz", memfile, freq, amp))

        results[filename] = entry_list

    document = Document()
    document.add_heading('Audio Quantization and Resampling Report', 0)

    for filename, entries in results.items():
        document.add_heading(f'File: {filename}', level=1)

        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Modification'
        hdr_cells[1].text = 'Max Frequency (Hz)'
        hdr_cells[2].text = 'Max Amplitude (dB)'

        for label, _, freq, amp in entries:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = f'{freq:.2f}'
            row_cells[2].text = f'{amp:.2f}'

        document.add_heading('Visualizations', level=2)
        for label, memfile, _, _ in entries:
            if memfile is not None:
                document.add_heading(label, level=3)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()

    document.save('quant_sampling_resampling_sound_report.docx')

# main()

def test_hearing_variants(file_path):
    print(f"test on: {file_path}")
    data, fs = sf.read(file_path)

    bit_depths = [4, 8]
    decimation_steps = [4, 6, 10, 24]
    interp_fs = [4000, 8000, 11999, 16000, 16953]
    interp_methods = ['linear', 'cubic']

    print("orig")
    sd.play(data, fs)
    sd.wait()

    for bits in bit_depths:
        print(f"quantization {bits}-bit")
        q_data = quantization(data, bits)
        sd.play(q_data, fs)
        sd.wait()
        time.sleep(1.5)

    for step in decimation_steps:
        print(f"decimation step {step}")
        d_data, new_fs = decimate(data, step, fs)
        sd.play(d_data, new_fs)
        sd.wait()
        time.sleep(1.5)

    for new_fs in interp_fs:
        for method in interp_methods:
            print(f"{method} interpolation {new_fs} Hz")
            try:
                i_data, _ = interpolate(data, fs, new_fs, method)
                sd.play(i_data, new_fs)
                sd.wait()
                time.sleep(1.5)
            except Exception as e:
                print(f"⚠️ Interpolation error: {e}")

    print(f"finished test on {file_path}")

test_files = [
    'SING/sing_low1.wav',
    'SING/sing_medium2.wav',
    'SING/sing_high1.wav'
]

for f in test_files:
    test_hearing_variants(f)


# data1, fs1 = sf.read('SIN/sin_8000Hz.wav')
# print(len(data) / fs)
# print(data[:111])