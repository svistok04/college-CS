import numpy as np
import sys
from docx import Document
from docx.shared import Inches
from io import BytesIO
import cv2
import matplotlib.pyplot as plt

def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

def rle_encode(data):
    shape = data.shape
    if len(shape) > 1:
        data = data.flatten()

    length = len(data)
    result = np.zeros(2 * length, np.uint8)
    cnt = 0
    cnt_result = 0
    for i in range(1, length):
        cnt += 1
        if data[i] != data[i-1] or cnt == 255:
            result[cnt_result] = cnt
            result[cnt_result + 1] = data[i - 1]
            cnt_result += 2
            cnt = 0

    result[cnt_result] = cnt + 1
    result[cnt_result + 1] = data[-1]
    cnt_result += 2

    return np.concatenate((
        np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.uint8),
                 result[:cnt_result]
    ))

def rle_decode(encoded):
    ndim = encoded[:4].view(np.uint32)[0]
    shape = encoded[4:4 + 4 * ndim].view(np.uint32)
    data_start = 4 + 4 * ndim

    data = encoded[data_start:]

    result = []
    for i in range(0, len(data), 2):
        count = data[i]
        value = data[i + 1]
        result.extend([value] * count)

    return np.array(result, dtype=np.uint8).reshape(tuple(shape))


def byterun_encode(data):
    shape = data.shape
    if len(shape) > 1:
        data = data.flatten()

    length = len(data)
    result = np.zeros(length * 2, dtype=np.uint8)
    out_idx = 0
    i = 0

    while i < length:
        run_start = i
        run_value = data[i]
        run_length = 1
        while i + run_length < length and data[i + run_length] == run_value and run_length < 127:
            # print(run_value)
            run_length += 1

        if run_length > 1:
            # print(f'repeated run length {run_length}')
            result[out_idx] = (-run_length + 256) % 256
            result[out_idx + 1] = run_value
            out_idx += 2
            i += run_length
        else:
            literal_start = i
            literal_length = 1
            while i + literal_length < length - 1 and (data[i + literal_length] != data[i + literal_length + 1] or literal_length == 1) and literal_length < 127:
                literal_length += 1
            result[out_idx] = literal_length
            result[out_idx + 1:out_idx + literal_length + 1] = data[i:i + literal_length]
            out_idx += 1 + literal_length
            i += literal_length

    return np.concatenate((
        np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.uint8),
                result[:out_idx].view(np.uint8)
    ))
    # return np.concatenate((
    #     np.array([len(shape)] + list(shape), dtype=np.uint32).view(np.int8),
    #     result[:out_idx]
    # ))

def byterun_decode(encoded: np.ndarray) -> np.ndarray:
    ndim = encoded[:4].view(np.uint32)[0]
    shape = encoded[4:4 + 4 * ndim].view(np.uint32)
    data_start = 4 + 4 * ndim

    data = encoded[data_start:].view(np.int8)
    result = []

    i = 0
    while i < len(data):
        count = data[i]
        if count < 0:
            value = data[i + 1]
            result.extend([value] * (-count))
            i += 2
        else:
            count = int(count)
            result.extend(data[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(result, dtype=np.uint8).reshape(tuple(shape))

def test_compression_methods():
    test_cases = [
        np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),
        np.array([1,2,3,1,2,3,1,2,3]),
        np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),
        np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]),
        np.zeros((1,520)),
        np.arange(0,521,1),
        np.eye(7),
        np.dstack([np.eye(7),np.eye(7),np.eye(7)]),
        np.ones((1,1,1,1,1,1,10))
    ]
    pass

def calculate_metrics(original, compressed):
    original_size = get_size(original)
    compressed_size = get_size(compressed)
    compression_ratio = round(original_size / compressed_size, 2)
    compression_percent = round((compressed_size / original_size) * 100, 2)
    return compression_ratio, compression_percent

def is_identical(original, decoded):
    return np.array_equal(original, decoded)

def process_image(img):
    img = img.astype(np.uint8)
    memfile = BytesIO()
    results = {}

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.imshow(img)
    ax.axis('off')
    plt.tight_layout()
    plt.savefig(memfile, format='png')
    plt.close(fig)
    memfile.seek(0)

    for name, encode, decode in [
        ("RLE", rle_encode, rle_decode),
        ("ByteRun", byterun_encode, byterun_decode)
    ]:
        encoded = encode(img)
        decoded = decode(encoded)
        correct = is_identical(img, decoded)
        ratio, percent = calculate_metrics(img, encoded)

        results[name] = {
            "correct": correct,
            "original_size": get_size(img),
            "compressed_size": get_size(encoded),
            "compression_ratio": ratio,
            "compression_percent": percent
        }

    return memfile, results

def main():
    filenames = ['drawing2.jpg', 'document1.jpg', 'color1.jpg', 'color2.jpg', 'color3.jpg']
    algos = ['rle', 'byterun']
    results = {}

    document = Document()
    document.add_heading('Lossless Compression Report', level=0)

    for filename in filenames:
        img = cv2.imread(filename, cv2.IMREAD_COLOR)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        memfile, results = process_image(img)

        document.add_page_break()
        document.add_heading(f"File: {filename}", level=1)
        document.add_picture(memfile, width=Inches(5.5))
        memfile.close()

        for name in ['RLE', 'ByteRun']:
            result = results[name]
            document.add_paragraph(f"{name}:")
            document.add_paragraph(f"Correctness: {'True' if result['correct'] else 'False'}")
            document.add_paragraph(f"Original size: {result['original_size']} bytes")
            document.add_paragraph(f"Compressed size: {result['compressed_size']} bytes")
            document.add_paragraph(f"Compression ratio: {result['compression_ratio']}")
            document.add_paragraph(f"Compression percent: {result['compression_percent']}%")
            document.add_paragraph("")

    document.save("compression_report.docx")



def test2():
    test1 = np.array([1,1,1,1,1,2,2,2,3,4,5,6,6,6,6,1])
    test2 = np.array([1,1,1,1,1,2,2,2,3,4,5,6,6,6,6,6])
    test3 = np.dstack([np.eye(7),np.eye(7),np.eye(7)])
    test4 = np.arange(0,20,1)

    print(test1)
    # test3_encode = rle_encode(test3)
    test1_encode = byterun_encode(test1)
    print(test1_encode[8:])
    print(byterun_decode(test1_encode))
    # print(f'{test3} -> {test3_encode}')
    # print(f'{test3_encode} -> {rle_decode(test3_encode)}')
    # print(f'{test2} -> {rle_encode(test2)}')
    # print(f'{test3} -> {rle_encode(test3)}')
    # print(f'{test4} -> {rle_encode(test4)}')

    # print(np.eye(7))

def test3():
    filenames = ['drawing2.jpg', 'document1.jpg', 'color1.jpg', 'color2.jpg', 'color3.jpg']
    for file in filenames:
        img = cv2.imread(file, cv2.IMREAD_COLOR)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        orig_size = get_size(img)
        print(img.shape)
        # plt.imshow(img)
        # plt.show()
        img = rle_encode(img)
        compressed_size = get_size(img)
        print(f'{file} {np.round(compressed_size/orig_size*100, 2)}% (of orig) compression')
        # img = rle_decode(img)
        # plt.imshow(img)
        # print(img.shape)
        # plt.show()

main()
# test2()
# test3()
