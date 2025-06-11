from time import time
from io import BytesIO
from docx import Document
from docx.shared import T, Inches
import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
from skimage.metrics import structural_similarity as ssim
# sprawzdic czy zaleznosc jest liniowa czy inna 


def read_gray(path):
    return cv2.imread(path, cv2.IMREAD_GRAYSCALE)

def degrade_jpeg(img, quality):
    encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), quality]
    _, encimg = cv2.imencode('.jpg', img, encode_param)
    return cv2.imdecode(encimg, 1)

def apply_gaussian_blur(img, ksize, sigma):
    return cv2.GaussianBlur(img, ksize, sigma)

def add_gaussian_noise(img, alpha, sigma):
    gauss = np.random.normal(0, sigma, img.shape)
    noisy = (img + alpha * gauss).clip(0, 255).astype(np.uint8)
    return noisy

def mse(K, I):
    return np.mean((I - K) ** 2)

def nmse(K, I):
    return np.sum((I - K) ** 2) / np.sum(K ** 2)

def psnr(K, I):
    mse_val = mse(K, I)
    return 10 * np.log10((255 ** 2) / mse_val) 

def image_fidelity(K, I):
    num = np.sum((K - I) ** 2)
    den = np.sum(K * I)
    return 1 - (num / den)

def ssim_score(K, I):
    return ssim(K, I)


def plot_metrics_separate(results, method_name, save_path=None):

    metrics = [
        ('MSE', 'lower is better'),
        ('NMSE', 'lower is better'),
        ('PSNR', 'higher is better'),
        ('IF', 'higher is better'),
        ('SSIM', 'higher is better'),
    ]

    levels = results['level']
    fig, axs = plt.subplots(2, 3, figsize=(12, 8))
    axs = axs.flatten() 

    for idx, (metric, note) in enumerate(metrics):
        axs[idx].plot(levels, results[metric], marker='o')
        axs[idx].set_title(f"{metric}\n({note})", fontsize=14)
        axs[idx].set_xlabel('JPEG Quality' if method_name.lower().startswith('jpeg') else 'Degradation Level')
        axs[idx].set_ylabel('Value')
        axs[idx].grid(True)
        axs[idx].set_xticks(levels)
        axs[idx].set_xticklabels(levels, rotation=45)

    if len(metrics) < len(axs):
        for ax in axs[len(metrics):]:
            ax.axis('off')

    fig.suptitle(f"{method_name.upper()} — Metric Trends", fontsize=22)
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150) 
    if save_path is not None:
        plt.savefig(save_path, dpi=150)
    plt.close()
    buf.seek(0)
    return buf

def plot_comparison(orig, degraded, metrics, method, level, save_path=None):
    fig, axs = plt.subplots(1, 3, figsize=(24, 8))
    axs[0].imshow(cv2.cvtColor(orig, cv2.COLOR_BGR2RGB))
    axs[0].set_title('Original')
    axs[1].imshow(cv2.cvtColor(degraded, cv2.COLOR_BGR2RGB))
    axs[1].set_title('Degraded')
    diff = cv2.absdiff(orig, degraded)
    axs[2].imshow(diff, cmap='hot')
    axs[2].set_title('Difference')
    metric_text = ', '.join([f"{k}={v:.2f}" for k,v in metrics.items()])
    fig.suptitle(f"{method}, Level={level} — {metric_text}", fontsize=28)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png')
    if save_path is not None:
        plt.savefig(save_path)
    plt.close()
    buf.seek(0)
    return buf

def run_degradation_experiment(image_path, method, param_list, sigma_blur=0, sigma_noise=25):
    orig_color = cv2.imread(image_path, cv2.IMREAD_COLOR)
    orig_gray = cv2.cvtColor(orig_color, cv2.COLOR_BGR2GRAY)
    degraded_images_color = []
    degraded_images_gray = []
    results = { 'level': [], 'MSE': [], 'NMSE': [], 'PSNR': [], 'IF': [], 'SSIM': [] }
    
    for param in param_list:
        if method == 'jpeg':
            degraded_color = degrade_jpeg(orig_color, param)
        elif method == 'blur':
            ksize = (param, param)
            degraded_color = cv2.GaussianBlur(orig_color, ksize, sigma_blur)
        elif method == 'noise':
            gauss = np.random.normal(0, sigma_noise, orig_color.shape)
            degraded_color = (orig_color + param * gauss).clip(0, 255).astype(np.uint8)
        else:
            raise ValueError("Unknown method")
        degraded_gray = cv2.cvtColor(degraded_color, cv2.COLOR_BGR2GRAY)
        degraded_images_color.append(degraded_color)
        degraded_images_gray.append(degraded_gray)
        m = mse(orig_gray, degraded_gray)
        n = nmse(orig_gray, degraded_gray)
        p = psnr(orig_gray, degraded_gray)
        i = image_fidelity(orig_gray, degraded_gray)
        s = ssim_score(orig_gray, degraded_gray)
        results['level'].append(param)
        results['MSE'].append(m)
        results['NMSE'].append(n)
        results['PSNR'].append(p)
        results['IF'].append(i)
        results['SSIM'].append(s)
    return orig_color, degraded_images_color, results

def add_metric_table(document, results, param_label):
    metrics = ['MSE', 'NMSE', 'PSNR', 'IF', 'SSIM']
    table = document.add_table(rows=1, cols=1 + len(metrics))
    table.style = 'Light List'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = param_label
    for i, m in enumerate(metrics):
        hdr_cells[i+1].text = m
    for i, level in enumerate(results['level']):
        row_cells = table.add_row().cells
        row_cells[0].text = str(level)
        for j, m in enumerate(metrics):
            row_cells[j+1].text = f"{results[m][i]:.4f}"

jpeg_levels = [75, 60, 45, 30, 15]
# jpeg_levels = [95, 30]
orig1, degs1, res1 = run_degradation_experiment("images/image1.jpg", 'jpeg', jpeg_levels)

blur_sigmas = [3,5,7,9,11]
# sigma_levels = [3, 11]
orig2, degs2, res2 = run_degradation_experiment("images/image2.jpg", 'blur',blur_sigmas, sigma_blur=0)

noise_alphas = [0.1, 0.25, 0.4, 0.6, 0.8]
# sigma_levels2 = [0.1, 0.8]
orig3, degs3, res3 = run_degradation_experiment("images/image3.jpg", 'noise', noise_alphas, sigma_noise=25)

document = Document()
document.add_heading("Image Quality Assessment – Objective Measures Report", 0)
document.add_page_break()

document.add_heading("Metrics Tables and Plots", level=1)
for idx, (orig, degs, res, label, levels) in enumerate([
    (orig1, degs1, res1, "JPEG Compression", jpeg_levels),
    (orig2, degs2, res2, "Gaussian Blur", blur_sigmas),
    (orig3, degs3, res3, "Gaussian Noise", noise_alphas),
]):
    document.add_heading(f"Image {idx+1} — {label}", level=2)
    add_metric_table(document, res, "Level")
    buf = plot_metrics_separate(res, label, save_path=f'results/metrics_{label}.png')
    document.add_picture(buf, width=Inches(6.5))

document.add_heading("Visual Comparison", level=1)
for idx, (orig, degs, res, label, levels) in enumerate([
    (orig1, degs1, res1, "JPEG Compression", jpeg_levels),
    (orig2, degs2, res2, "Gaussian Blur", blur_sigmas),
    (orig3, degs3, res3, "Gaussian Noise", noise_alphas),
]):
    document.add_heading(f"Image {idx+1} — {label}", level=2)
    for j in range(len(degs)):
        metrics = {k: res[k][j] for k in ['MSE','NMSE','PSNR','IF','SSIM']}
        buf = plot_comparison(orig, degs[j], metrics, label, levels[j], save_path=f"results/compare_{label}_{j}.png")
        document.add_picture(buf, width=Inches(5.5))

document.save('image_quality_lab.docx')
